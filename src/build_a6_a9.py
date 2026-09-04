from __future__ import annotations

from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from course_common import (
    NAVY, TEAL, TEAL_DARK, PALE, PALE_TEAL, WARM, MID, TEXT, WHITE,
    base_doc, block, add_text, add_step, add_tip, add_check, add_finish,
    add_picture, new_workspace_section, finalise, set_run, _clear, _keep_first,
    _fill, _border,
)


def fonts():
    root = Path('/usr/share/fonts/truetype/liberation2')
    return str(root/'LiberationSans-Regular.ttf'), str(root/'LiberationSans-Bold.ttf')


def school_icon(path: Path):
    W,H=420,300
    im=Image.new('RGBA',(W,H),(255,255,255,0)); d=ImageDraw.Draw(im)
    d.rounded_rectangle((55,95,365,260),radius=12,fill='#EAF4F3',outline='#237B78',width=5)
    d.polygon([(45,105),(210,35),(375,105)],fill='#237B78')
    d.rounded_rectangle((180,170,240,260),radius=6,fill='#17324D')
    for x in (95,275): d.rounded_rectangle((x,145,x+55,195),radius=5,fill='white',outline='#17324D',width=4)
    d.line((210,35,210,2),fill='#17324D',width=5); d.polygon([(210,4),(274,20),(210,36)],fill='#17324D')
    im.save(path)


def a7_source(path: Path, icon_path: Path):
    icon=Image.open(icon_path).convert('RGBA')
    canvas=Image.new('RGB',(1100,760),'white')
    icon.thumbnail((650,470)); x=(1100-icon.width)//2; y=(760-icon.height)//2
    canvas.paste(icon.convert('RGB'),(x,y)); d=ImageDraw.Draw(canvas); d.rectangle((1,1,1098,758),outline='#D3DEE2',width=5)
    canvas.save(path,quality=95)


def a7_preview(path: Path, icon_path: Path):
    reg,bold=fonts(); W,H=1500,480
    im=Image.new('RGB',(W,H),'white'); d=ImageDraw.Draw(im)
    ft=ImageFont.truetype(bold,46); fb=ImageFont.truetype(reg,27); fs=ImageFont.truetype(bold,24)
    d.rounded_rectangle((18,18,W-18,H-18),radius=18,outline='#D3DEE2',width=3,fill='white')
    d.text((70,48),'UNSER SCHULHAUS',font=ft,fill='#17324D'); d.line((70,112,1430,112),fill='#237B78',width=4)
    icon=Image.open(icon_path).convert('RGBA'); icon.thumbnail((410,285)); im.paste(icon,(990,145),icon)
    y=155
    for s in ['Im Schulhaus Sonnenberg lernen rund 240 Schülerinnen','und Schüler.','','Die Bibliothek befindet sich im Erdgeschoss.','In der grossen Pause ist sie geöffnet.']:
        if s: d.text((75,y),s,font=fb,fill='#17324D')
        y+=46
    d.text((75,394),'Das Bild steht rechts. Der Text läuft links daran vorbei.',font=fs,fill='#5E6D78')
    im.save(path,quality=95)


def a8_preview(path: Path):
    reg,bold=fonts(); W,H=1500,500
    im=Image.new('RGB',(W,H),'white'); d=ImageDraw.Draw(im)
    ft=ImageFont.truetype(bold,32); fh=ImageFont.truetype(bold,23); fb=ImageFont.truetype(reg,22)
    d.rounded_rectangle((20,20,W-20,H-20),radius=18,outline='#D3DEE2',width=3,fill='white')
    x0,y0=70,55; tw=1360; cols=[230,330,520,280]; line='#9FB0B7'
    d.rectangle((x0,y0,x0+tw,y0+60),fill='white',outline=line,width=2)
    s='SPORTTAG – ABLAUF'; box=d.textbbox((0,0),s,font=ft); d.text((x0+(tw-(box[2]-box[0]))/2,y0+11),s,font=ft,fill='#17324D')
    rows=[['Zeit','Ort','Aktivität','Gruppe'],['09.00','Aula','Begrüssung','Alle'],['09.30','Turnhalle','Staffellauf','A'],['09.30','Sportplatz','Fussball','B'],['10.30','Aula','Pause','Alle']]
    y=y0+60
    for ri,row in enumerate(rows):
        x=x0
        for j,(w,val) in enumerate(zip(cols,row)):
            d.rectangle((x,y,x+w,y+54),fill='white',outline=line,width=2); f=fh if ri==0 else fb; c='#237B78' if ri==0 else '#17324D'; box=d.textbbox((0,0),val,font=f); tx=x+(w-(box[2]-box[0]))/2 if ri==0 or j in (0,3) else x+14; d.text((tx,y+13),val,font=f,fill=c); x+=w
        y+=54
    im.save(path,quality=95)


def a9_preview(path: Path):
    reg,bold=fonts(); W,H=1500,520
    im=Image.new('RGB',(W,H),'white'); d=ImageDraw.Draw(im)
    ft=ImageFont.truetype(bold,46); f1=ImageFont.truetype(bold,30); f2=ImageFont.truetype(bold,24); fb=ImageFont.truetype(reg,22)
    d.rounded_rectangle((18,18,W-18,H-18),radius=18,outline='#D3DEE2',width=3,fill='white')
    d.text((65,40),'UNSER SCHULAUSFLUG NACH LUZERN',font=ft,fill='#17324D')
    d.text((65,125),'Der Morgen',font=f1,fill='#237B78'); d.text((65,170),'Treffpunkt',font=f2,fill='#17324D'); d.text((65,205),'Wir treffen uns um 08.00 Uhr beim Schulhaus und fahren gemeinsam zum Bahnhof.',font=fb,fill='#5E6D78')
    d.text((65,270),'Im Verkehrshaus',font=f1,fill='#237B78'); d.text((65,315),'Unsere Aufgabe',font=f2,fill='#17324D'); d.text((65,350),'In Gruppen suchen wir drei Ausstellungsstücke und notieren die wichtigsten Informationen.',font=fb,fill='#5E6D78')
    d.text((65,415),'Mittag und Altstadt',font=f1,fill='#237B78'); d.text((65,458),'Nach dem Mittagessen spazieren wir gemeinsam durch die Altstadt.',font=fb,fill='#5E6D78')
    im.save(path,quality=95)


def build_a6(out: Path):
    doc=base_doc('A6','Seitenlayout','Die Seite passend einstellen','Du arbeitest heute nur am Aufbau der Seite – nicht am Text.','eine Seite auf Querformat stellen, Seitenränder ändern und mit einem Seitenumbruch gezielt eine neue Seite beginnen.')
    t=block(doc,'01','SEITE 2'); r=t.cell(0,1); p=r.paragraphs[0]; _clear(p); x=p.add_run('Bearbeite nur die Übungsseite'); set_run(x,size=12.8,bold=True,color=NAVY); add_text(r,'Klicke zuerst irgendwo auf Seite 2. Der Text dort ist bereits richtig formatiert.',9.5,after=.7)
    add_step(r,'A',[('Seite 2',True,False,NAVY),('  →  Querformat',False,False,None)])
    add_step(r,'B',[('Seite 2',True,False,NAVY),('  →  Seitenränder «Schmal»',False,False,None)])
    add_step(r,'C',[('Klicke direkt vor «PACKLISTE»',True,False,NAVY),('  →  Ctrl + Enter',False,False,None)])
    add_step(r,'D',[('Kontrolliere',True,False,NAVY),('  →  «PACKLISTE» beginnt jetzt auf einer neuen Seite.',False,False,None)])
    add_tip(doc,'Hochformat = ▯   |   Querformat = ▭   |   Neue Seite = Ctrl + Enter','MERKE')
    add_tip(doc,'Für eine neue Seite nicht viele Male Enter drücken. Ein Seitenumbruch bleibt auch dann richtig, wenn später Text dazukommt.')
    add_check(doc,'Seite 1 bleibt Hochformat. Die Übungsseite ist Querformat. «PACKLISTE» beginnt auf einer neuen Seite.')
    add_finish(doc)
    sec=new_workspace_section(doc,'A6'); sec.orientation=WD_ORIENT.PORTRAIT; sec.page_width=Cm(21); sec.page_height=Cm(29.7)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; x=p.add_run('SCHULAUSFLUG NACH LUZERN'); set_run(x,size=20,bold=True,color=NAVY)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(12); x=p.add_run('Dienstag, 22. September'); set_run(x,size=11,bold=True,color=TEAL)
    p=doc.add_paragraph(); x=p.add_run('PROGRAMM'); set_run(x,size=13,bold=True,color=NAVY)
    for tm,txt in [('08.00 Uhr','Treffpunkt beim Schulhaus'),('08.20 Uhr','Abfahrt mit dem Zug'),('10.00 Uhr','Verkehrshaus'),('12.15 Uhr','Mittagspause am See'),('14.00 Uhr','Altstadt-Rundgang'),('16.30 Uhr','Rückfahrt')]:
        p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(4); x=p.add_run(tm+'  –  '); set_run(x,size=11,bold=True,color=NAVY); x=p.add_run(txt); set_run(x,size=11)
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(14); x=p.add_run('PACKLISTE'); set_run(x,size=13,bold=True,color=NAVY)
    for item in ['Trinkflasche','Lunch','Regenjacke','Schreibzeug','Halbtax / Billet falls vorhanden']:
        p=doc.add_paragraph(); p.paragraph_format.left_indent=Cm(.5); x=p.add_run('•  '+item); set_run(x,size=11)
    finalise(doc,out,'A6 - Seitenlayout')


def build_a7(out: Path, preview: Path):
    doc=base_doc('A7','Bilder in Word','Ein Bild passend einsetzen','Du fügst ein Bild ein und passt es so an, dass Text und Bild zusammenpassen.','ein Bild aus einer Datei einfügen, zuschneiden, auf eine passende Grösse bringen und den Text darum laufen lassen.')
    t=block(doc,'01','AUFGABE'); r=t.cell(0,1); p=r.paragraphs[0]; _clear(p); x=p.add_run('Baue das Infoblatt nach'); set_run(x,size=12.8,bold=True,color=NAVY); add_text(r,'Bilddatei: a7_schulhaus.png',9.5,bold=True,color=TEAL_DARK); add_picture(r,preview,11.1)
    for letter,parts in [
        ('A',[('Klicke auf «[BILD HIER EINFÜGEN]»',True,False,NAVY),('  →  Bild a7_schulhaus.png einfügen',False,False,None)]),
        ('B',[('Bild zuschneiden',True,False,NAVY),('  →  den grossen weissen Rand entfernen',False,False,None)]),
        ('C',[('Bildbreite',True,False,NAVY),('  →  ungefähr 5,5 cm',False,False,None)]),
        ('D',[('Textumbruch',True,False,NAVY),('  →  Quadrat',False,False,None)]),
        ('E',[('Bild',True,False,NAVY),('  →  rechts neben den Text verschieben',False,False,None)]),
        ('F',[('«[BILD HIER EINFÜGEN]»',True,False,NAVY),('  →  löschen',False,False,None)]),
    ]: add_step(r,letter,parts)
    t=block(doc,'HIER','ARBEITEN',fill_left=PALE,fill_right='FBFCFC',label_color=TEAL_DARK,label_size=11.0); r=t.cell(0,1); p=r.paragraphs[0]; _clear(p); x=p.add_run('UNSER SCHULHAUS'); set_run(x,size=18,bold=True,color=NAVY); add_text(r,'[BILD HIER EINFÜGEN]',9.5,bold=True,color=TEAL_DARK); add_text(r,'Im Schulhaus Sonnenberg lernen rund 240 Schülerinnen und Schüler.',10.6); add_text(r,'Die Bibliothek befindet sich im Erdgeschoss.',10.6); add_text(r,'In der grossen Pause ist sie geöffnet.',10.6)
    add_tip(doc,'Bild anklicken → Bildformat. Dort findest du Zuschneiden. Für den Textumbruch klickst du auf das kleine Layout-Symbol neben dem Bild.')
    add_check(doc,'Ist der weisse Rand weg? Steht das Bild rechts? Läuft der Text links am Bild vorbei?')
    add_finish(doc,'Gib dieses Arbeitsblatt zusammen mit der Bilddatei in deinem Ordner "IB" ab.')
    finalise(doc,out,'A7 - Bilder in Word')


def build_a8(out: Path, preview: Path):
    doc=base_doc('A8','Tabellen','Infos ins Raster bringen','Du baust eine einfache Tabelle und passt sie Schritt für Schritt an.','eine Tabelle erstellen, Daten in Zellen eintragen, eine Zeile ergänzen und Zellen verbinden.')
    t=block(doc,'01','AUFGABE'); r=t.cell(0,1); p=r.paragraphs[0]; _clear(p); x=p.add_run('Baue den Sporttag-Plan nach'); set_run(x,size=12.8,bold=True,color=NAVY); add_text(r,'Arbeite auf Seite 2. Die Daten stehen dort bereits bereit.',9.5); add_picture(r,preview,10.9)
    for letter,parts in [
        ('A',[('Unter «HIER TABELLE EINFÜGEN»',True,False,NAVY),('  →  Tabelle mit 4 Spalten und 5 Zeilen einfügen',False,False,None)]),
        ('B',[('Kopfzeile + vier Datenzeilen',True,False,NAVY),('  →  Angaben von oben in die richtigen Zellen übertragen',False,False,None)]),
        ('C',[('Kopfzeile',True,False,NAVY),('  →  fett und zentriert',False,False,None)]),
        ('D',[('Zeit und Gruppe',True,False,NAVY),('  →  zentrieren',False,False,None)]),
        ('E',[('Neue Zeile über der Tabelle einfügen',True,False,NAVY),('  →  alle 4 Zellen dieser Zeile verbinden',False,False,None)]),
        ('F',[('In die verbundene Zelle',True,False,NAVY),('  →  «SPORTTAG – ABLAUF», 16 pt, fett, zentriert',False,False,None)]),
    ]: add_step(r,letter,parts)
    add_tip(doc,'Zellen verbinden: Markiere die vier Zellen der neuen Zeile. Unter Tabellenlayout findest du «Zellen verbinden».')
    add_check(doc,'Hat deine Tabelle 4 Spalten? Sind alle Daten in der richtigen Zeile? Geht der Titel über die ganze Tabelle?')
    add_finish(doc)
    new_workspace_section(doc,'A8')
    p=doc.add_paragraph(); x=p.add_run('DATEN FÜR DIE TABELLE'); set_run(x,size=16,bold=True,color=NAVY)
    for s in ['Zeit | Ort | Aktivität | Gruppe','09.00 | Aula | Begrüssung | Alle','09.30 | Turnhalle | Staffellauf | A','09.30 | Sportplatz | Fussball | B','10.30 | Aula | Pause | Alle']:
        p=doc.add_paragraph(); x=p.add_run(s); set_run(x,size=11)
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(12); x=p.add_run('HIER TABELLE EINFÜGEN'); set_run(x,size=12,bold=True,color=TEAL_DARK)
    for _ in range(12): doc.add_paragraph(' ')
    finalise(doc,out,'A8 - Tabellen')


def build_a9(out: Path, preview: Path):
    doc=base_doc('A9','Formatvorlagen','Überschriften mit System','Du gibst Textteilen eine feste Rolle, statt jede Überschrift von Hand zu gestalten.','Titel, Überschrift 1, Überschrift 2 und normalen Text mit den passenden Word-Formatvorlagen auszeichnen.')
    t=block(doc,'01','SEITE 2'); r=t.cell(0,1); p=r.paragraphs[0]; _clear(p); x=p.add_run('Gib jedem Textteil die richtige Rolle'); set_run(x,size=12.8,bold=True,color=NAVY); add_text(r,'Arbeite auf Seite 2. Ändere Schriftgrösse oder Fett nicht von Hand.',9.5); add_picture(r,preview,10.9)
    add_step(r,'A',[('«UNSER SCHULAUSFLUG NACH LUZERN»',True,False,NAVY),('  →  Formatvorlage «Titel»',False,False,None)])
    add_step(r,'B',[('«Der Morgen», «Im Verkehrshaus», «Mittag und Altstadt»',True,False,NAVY),('  →  «Überschrift 1»',False,False,None)])
    add_step(r,'C',[('«Treffpunkt», «Unsere Aufgabe», «Rückfahrt»',True,False,NAVY),('  →  «Überschrift 2»',False,False,None)])
    add_step(r,'D',[('Alle übrigen Absätze',True,False,NAVY),('  →  «Standard»',False,False,None)])
    add_tip(doc,'Titel  >  Überschrift 1  >  Überschrift 2  >  Standard','MERKE')
    add_tip(doc,'Formatvorlagen findest du bei «Start». Klicke zuerst in den Absatz und dann auf die passende Formatvorlage.')
    add_check(doc,'Sehen alle Überschriften derselben Stufe gleich aus? Dann hast du die Formatvorlagen richtig eingesetzt.')
    add_finish(doc)
    new_workspace_section(doc,'A9')
    raw=[
        'UNSER SCHULAUSFLUG NACH LUZERN','Der Morgen','Treffpunkt','Wir treffen uns um 08.00 Uhr beim Schulhaus und fahren gemeinsam zum Bahnhof. Im Zug sitzen wir in unseren Gruppen.','Im Verkehrshaus','Unsere Aufgabe','In Gruppen suchen wir drei Ausstellungsstücke. Zu jedem Stück notieren wir zwei wichtige Informationen.','Mittag und Altstadt','Nach dem Mittagessen spazieren wir gemeinsam durch die Altstadt. Danach bleibt noch Zeit für eine kurze Pause am See.','Rückfahrt','Um 16.30 Uhr fahren wir zurück. Die Ankunft beim Schulhaus ist ungefähr um 18.00 Uhr.'
    ]
    for s in raw:
        p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(6); x=p.add_run(s); set_run(x,size=11)
    finalise(doc,out,'A9 - Formatvorlagen und Überschriften')


def build_all(root: Path):
    sheets=root/'arbeitsblaetter'; assets=sheets/'assets'; prev=assets/'vorlagen'; assets.mkdir(parents=True,exist_ok=True); prev.mkdir(parents=True,exist_ok=True)
    icon=assets/'a7_school_icon.png'; school_icon(icon)
    a7src=assets/'a7_schulhaus.png'; a7prev=prev/'a7_schulhaus_vorlage.png'; a8prev=prev/'a8_sporttag_tabelle_vorlage.png'; a9prev=prev/'a9_formatvorlagen_vorlage.png'
    a7_source(a7src,icon); a7_preview(a7prev,icon); a8_preview(a8prev); a9_preview(a9prev)
    build_a6(sheets/'A6_Seitenlayout.docx')
    build_a7(sheets/'A7_Bilder_in_Word.docx',a7prev)
    build_a8(sheets/'A8_Tabellen.docx',a8prev)
    build_a9(sheets/'A9_Formatvorlagen_und_Ueberschriften.docx',a9prev)
