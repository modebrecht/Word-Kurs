from __future__ import annotations

from pathlib import Path

from PIL import Image, ImageDraw
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION

from course_common import (
    NAVY, TEAL, TEAL_DARK, PALE, PALE_TEAL, WARM, MID, LIGHT_LINE,
    WHITE, TEXT, base_doc, block, add_text, add_finish, finalise, set_run,
)
from course_build_helpers import (
    clear_paragraph as _clear,
    style_run as _font,
    fill_cell as _fill,
    set_cell_margins as _margins,
    set_cell_borders as _border,
)
from grading import swiss_grade_str


def test_asset(path: Path):
    W, H = 1500, 900
    core = Image.new("RGB", (W, H), "#DDEBF2")
    d = ImageDraw.Draw(core)
    d.rectangle((0,0,W,390), fill="#DCEBF3")
    d.polygon([(0,390),(300,280),(520,370),(850,245),(1120,360),(1500,300),(1500,560),(0,560)], fill="#6E8A7D")
    d.rectangle((0,500,W,H), fill="#79B8C3")
    d.polygon([(560,310),(900,280),(1010,740),(420,740)], fill="#D9F0F2")
    for x in range(485,990,70):
        d.line((x,340,x-45,720), fill="#FFFFFF", width=18)
    d.ellipse((290,680,1150,850), fill="#E9F6F6")
    for pts in [
        [(80,650),(270,500),(420,690),(300,800)],
        [(1120,640),(1320,475),(1490,690),(1430,830)],
        [(720,610),(810,500),(930,650),(850,760)],
    ]:
        d.polygon(pts, fill="#5E6F70", outline="#17324D")
    for x,y,s in [(150,420,1.0),(1320,390,1.05),(1180,440,.8)]:
        d.rectangle((x-10,y+35,x+10,y+150), fill="#755B45")
        d.ellipse((x-85*s,y-80*s,x+85*s,y+75*s), fill="#456C5F")
    d.line((1030,330,1030,470), fill="#17324D", width=5)
    d.polygon([(1030,332),(1115,350),(1030,380)], fill="#D95852")
    d.rectangle((4,4,W-5,H-5), outline="#D3DEE2", width=8)
    border = 140
    img = Image.new("RGB", (W + 2*border, H + 2*border), "white")
    img.paste(core, (border,border))
    img.save(path, quality=95)


def _add_req(cell, no, title, text, points):
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(.32)
    p.paragraph_format.line_spacing = 1
    r = p.add_run(f"{no}. ")
    set_run(r, size=9.0, bold=True, color=TEAL_DARK)
    r = p.add_run(title + " - ")
    set_run(r, size=9.0, bold=True, color=NAVY)
    r = p.add_run(text)
    set_run(r, size=9.0, color=TEXT)
    r = p.add_run(f"  [{points} P]")
    set_run(r, size=8.7, bold=True, color=MID)


def _test_header_footer(doc, correction=False):
    sec = doc.sections[0]
    ht = sec.header.tables[0]
    ht.cell(0,2).text = "KORREKTUR" if correction else "WORD-TEST"
    for j, c in enumerate(ht.rows[0].cells):
        for p in c.paragraphs:
            for r in p.runs:
                _font(r, size=9.2, bold=True, color=WHITE if j == 0 else (TEAL_DARK if j == 2 else MID))
    ft = sec.footer.tables[0]
    ft.cell(0,1).text = "WORD-TEST · KORREKTUR" if correction else "BENOTETER WORD-TEST"
    for c in ft.rows[0].cells:
        for p in c.paragraphs:
            for r in p.runs:
                _font(r, size=8.2, bold=True, color=MID)


def build_task_sheet(out: Path):
    doc = base_doc(
        "TEST",
        "Word-Test",
        "Schulreise zum Rheinfall",
        "Arbeitszeit: 45 Minuten · maximal 30 Punkte",
        "bekannte Word-Funktionen unter Prüfungsbedingungen selbstständig anwenden.",
    )
    _test_header_footer(doc)

    t = block(doc, "START", None, fill_left=NAVY, fill_right=PALE, label_color=WHITE, label_size=9.2)
    r = t.cell(0,1)
    p = r.paragraphs[0]
    _clear(p)
    x = p.add_run("Öffne ")
    set_run(x, size=9.55, bold=True, color=NAVY)
    x = p.add_run("Word_Test_Ausgangsdokument.docx")
    set_run(x, size=9.55, bold=True, color=TEAL_DARK)
    x = p.add_run(" und speichere sofort eine Kopie als ")
    set_run(x, size=9.55, color=TEXT)
    x = p.add_run("WordTest_Nachname_Vorname.docx")
    set_run(x, size=9.55, bold=True, color=NAVY)
    add_text(r, "Arbeite nur in deiner Kopie. Die Bilddatei heisst word_test_rheinfall.png.", 9.2, after=0)

    t = block(doc, "01-05", "TEXT & SEITE")
    r = t.cell(0,1)
    _clear(r.paragraphs[0])
    _add_req(r,1,"Datei","Kopie mit dem verlangten Dateinamen speichern.",1)
    _add_req(r,2,"Seite","A4 Hochformat; Seitenränder «Schmal».",2)
    _add_req(r,3,"Grundtext","Ganzer Text Arial 11 pt; Zeilenabstand 1,15; Absatzabstand danach 6 pt.",3)
    _add_req(r,4,"Formatvorlagen","Dokumenttitel = «Titel»; ÜBER DIE SCHULREISE, MITNEHMEN, ABLAUF, PROGRAMM und WICHTIG = «Überschrift 1».",4)
    _add_req(r,5,"Hervorhebung","Die Zeile «Freitag, 21. Mai 2027» fett und dunkelblau formatieren.",2)

    t = block(doc, "06-08", "LISTEN & BILD", fill_left=WARM, fill_right=WARM, label_color=NAVY)
    r = t.cell(0,1)
    _clear(r.paragraphs[0])
    _add_req(r,6,"Aufzählung","Die fünf Einträge unter MITNEHMEN als echte Word-Aufzählung formatieren.",2)
    _add_req(r,7,"Nummerierung","Die vier Schritte unter ABLAUF als echte Word-Nummerierung formatieren.",2)
    _add_req(r,8,"Bild","Ersetze die Zeile «BILD HIER EINFÜGEN: word_test_rheinfall.png» durch das Bild; weissen Rand wegschneiden; ca. 6 cm breit; Textumbruch «Quadrat»; rechts neben dem Einleitungstext platzieren.",4)

    t = block(doc, "09-11", "SEITE 2", fill_left=PALE_TEAL, fill_right=PALE_TEAL, label_color=TEAL_DARK)
    r = t.cell(0,1)
    _clear(r.paragraphs[0])
    _add_req(r,9,"Seitenumbruch","Direkt vor PROGRAMM mit Ctrl + Enter eine neue Seite beginnen.",2)
    _add_req(r,10,"Tabelle","Programmdaten als echte Tabelle mit 4 Spalten und 5 Zeilen übertragen; Kopfzeile fett und zentriert; Spalte «Zeit» zentrieren.",5)
    _add_req(r,11,"Kopf/Fuss","Kopfzeile «SCHULREISE RHEINFALL 2027»; unten mittig «Seite » plus automatische Seitenzahl.",3)

    t = block(doc, "CHECK", None, fill_left=WHITE, fill_right=WHITE, label_color=TEAL_DARK, label_size=9.2)
    r = t.cell(0,1)
    p = r.paragraphs[0]
    _clear(p)
    x = p.add_run("Vor dem Abgeben: ")
    set_run(x, size=9.4, bold=True, color=NAVY)
    x = p.add_run("genau 2 Seiten · Bild nicht verzerrt · echte Listen/Tabelle · Kopfzeile auf beiden Seiten · Seitenzahlen automatisch · Dateiname korrekt.")
    set_run(x, size=9.15, color=TEXT)
    add_finish(doc, "Speichere ein letztes Mal und gib deine DOCX-Datei im vorgesehenen Schulordner ab.")
    finalise(doc, out, "Benoteter Word-Test")


def build_source_document(out: Path):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(10.5)
    raw = [
        ("SCHULREISE ZUM RHEINFALL",4),("Freitag, 21. Mai 2027",5),("ÜBER DIE SCHULREISE",3),
        ("Unsere Klasse reist mit dem Zug nach Neuhausen am Rheinfall. Dort erkunden wir den Aussichtspunkt und arbeiten in kleinen Gruppen an einem kurzen Beobachtungsauftrag.",4),
        ("Nach dem Mittag gehen wir gemeinsam zum Uferweg. Am Nachmittag bleibt Zeit für einen Rundgang, bevor wir zum Bahnhof zurückkehren.",5),
        ("BILD HIER EINFÜGEN: word_test_rheinfall.png",5),("MITNEHMEN",3),
        ("Trinkflasche",1),("Lunch",1),("Regenjacke",1),("Schreibzeug",1),("kleiner Rucksack",5),
        ("ABLAUF",3),("Treffpunkt beim Schulhaus",1),("Zugfahrt nach Neuhausen",1),("Gruppenauftrag beim Rheinfall",1),("Rückfahrt zum Schulhaus",5),
        ("PROGRAMM",3),("Zeit | Ort | Aktivität | Hinweis",1),("08.00 | Schulhaus | Treffpunkt | pünktlich sein",1),
        ("09.30 | Rheinfall | Beobachtungsauftrag | in Gruppen",1),("12.15 | Picknickplatz | Mittagessen | eigener Lunch",1),
        ("15.45 | Bahnhof | Rückfahrt | gemeinsam",5),("WICHTIG",3),
        ("Wir bleiben während der Schulreise in den eingeteilten Gruppen. Treffpunkte und Zeiten müssen eingehalten werden. Am Wasser beachten wir die Anweisungen der Lehrperson.",3),
    ]
    for text, after in raw:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(after)
        p.paragraph_format.line_spacing = 1
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(10.5)
    doc.core_properties.title = "Word-Test Ausgangsdokument"
    doc.core_properties.subject = "Word Kurs Sek 8"
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)


def build_correction_sheet(out: Path):
    doc = base_doc(
        "TEST", "Korrektur", "Korrekturblatt",
        "30 Punkte · lineare Schweizer Notenskala · 18 Punkte = Note 4.0",
        "den Word-Test mit klaren Teilpunkten schnell und konsistent korrigieren.",
    )
    _test_header_footer(doc, correction=True)
    criteria = [
        ("1","Dateiname korrekt",1),("2","A4 Hochformat + Seitenränder schmal",2),
        ("3","Arial 11; 1,15 Zeilenabstand; 6 pt danach",3),("4","Titel + fünf Bereiche mit verlangten Formatvorlagen",4),
        ("5","Datumszeile fett + dunkelblau",2),("6","MITNEHMEN als echte Aufzählung",2),
        ("7","ABLAUF als echte Nummerierung",2),("8","Bild: Platzhalter ersetzt, zugeschnitten, ca. 6 cm, Quadrat, rechts",4),
        ("9","Echter Seitenumbruch vor PROGRAMM",2),("10","Tabelle 4×5; Kopfzeile fett/zentriert; Zeit zentriert",5),
        ("11","Kopfzeile korrekt + automatische Seitenzahl unten mittig",3),
    ]
    score_details = {
        "8": "Teilpunkte: Bild statt Platzhalter 1 · Zuschneiden 1 · Grösse 0,5 · Quadrat 0,5 · Position 1",
        "10": "Teilpunkte: echte 4×5-Tabelle 2 · Daten korrekt 1 · Kopfzeile 1 · Zeit zentriert 1",
        "11": "Teilpunkte: Kopfzeile 1 · automatische Seitenzahl 1 · Fusszeile mittig/auf beiden Seiten 1",
    }
    t = block(doc, "PUNKTE", None, fill_left=NAVY, fill_right=PALE, label_color=WHITE, label_size=9.2)
    r = t.cell(0,1)
    p = r.paragraphs[0]
    _clear(p)
    x = p.add_run("Name: ______________________________________     Punkte: ______ / 30     Note: ______")
    set_run(x, size=10.1, bold=True, color=NAVY)

    groups = [criteria[:4], criteria[4:8], criteria[8:]]
    labels = [("01-04","GRUNDLAGEN",TEAL,WHITE),("05-08","FORMAT & BILD",WARM,NAVY),("09-11","ABSCHLUSS",PALE_TEAL,TEAL_DARK)]
    for group, (left, sub, fill, color) in zip(groups, labels):
        t = block(doc, left, sub, fill_left=fill, fill_right=fill if fill != TEAL else WHITE, label_color=color, label_size=12.8)
        r = t.cell(0,1)
        _clear(r.paragraphs[0])
        for no, crit, pts in group:
            p = r.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(.45)
            x = p.add_run(no + ". ")
            set_run(x, size=9.2, bold=True, color=TEAL_DARK)
            x = p.add_run(crit)
            set_run(x, size=9.2, color=TEXT)
            x = p.add_run(f"    ____ / {pts}")
            set_run(x, size=9.2, bold=True, color=NAVY)
            detail = score_details.get(no)
            if detail:
                x = p.add_run("\n" + detail)
                set_run(x, size=8.0, color=MID)

    t = block(doc, "NOTE", None, fill_left=WHITE, fill_right=WHITE, label_color=TEAL_DARK, label_size=9.2)
    r = t.cell(0,1)
    p = r.paragraphs[0]
    _clear(p)
    x = p.add_run("Formel: Note = 1 + 5 × (Punkte / 30). ")
    set_run(x, size=9.5, bold=True, color=NAVY)
    x = p.add_run("Kaufmännisch auf eine Dezimalstelle runden. 18/30 = 60 % = Note 4.0.")
    set_run(x, size=9.3, color=TEXT)

    sec2 = doc.add_section(WD_SECTION.NEW_PAGE)
    sec2.top_margin = Cm(1.7)
    sec2.bottom_margin = Cm(1.5)
    sec2.left_margin = Cm(2.0)
    sec2.right_margin = Cm(2.0)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    x = p.add_run("NOTENSCHLÜSSEL")
    set_run(x, size=10, bold=True, color=TEAL_DARK)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    x = p.add_run("30 Punkte → Schweizer Note")
    set_run(x, size=25, bold=True, color=NAVY)

    mapping = [(pts, swiss_grade_str(pts, 30)) for pts in range(30,-1,-1)]
    tbl = doc.add_table(rows=17, cols=4)
    tbl.autofit = False
    widths = [Cm(3.0)] * 4
    for i, c in enumerate(tbl.rows[0].cells):
        c.width = widths[i]
        _fill(c, PALE_TEAL)
        _margins(c,80,100,80,100)
        p = c.paragraphs[0]
        _clear(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        x = p.add_run(["Punkte","Note","Punkte","Note"][i])
        set_run(x,size=9.6,bold=True,color=NAVY)
    left = mapping[:16]
    right = mapping[16:]
    for ri in range(16):
        vals = [str(left[ri][0]), left[ri][1], str(right[ri][0]) if ri < len(right) else "", right[ri][1] if ri < len(right) else ""]
        row = tbl.rows[ri+1]
        for ci, (c, val) in enumerate(zip(row.cells, vals)):
            c.width = widths[ci]
            _margins(c,65,90,65,90)
            p = c.paragraphs[0]
            _clear(p)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            x = p.add_run(val)
            set_run(x,size=9.2,bold=(ci in (1,3)),color=NAVY if ci in (1,3) else TEXT)
            _border(c,top={"val":"single","sz":"4","color":LIGHT_LINE},bottom={"val":"single","sz":"4","color":LIGHT_LINE},left={"val":"single","sz":"4","color":LIGHT_LINE},right={"val":"single","sz":"4","color":LIGHT_LINE})
        if left[ri][0] == 18:
            _fill(row.cells[0], WARM)
            _fill(row.cells[1], WARM)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    x = p.add_run("Bestehensgrenze: 18 von 30 Punkten = 60 % = Note 4.0.")
    set_run(x, size=10.2, bold=True, color=TEAL_DARK)
    finalise(doc, out, "Word-Test Korrekturblatt")


def build_all(root: Path):
    sheets = root / "arbeitsblaetter"
    assets = sheets / "assets"
    assets.mkdir(parents=True, exist_ok=True)
    test_asset(assets / "word_test_rheinfall.png")
    build_task_sheet(sheets / "Word_Test.docx")
    build_source_document(sheets / "Word_Test_Ausgangsdokument.docx")
    build_correction_sheet(sheets / "Word_Test_Korrekturblatt.docx")