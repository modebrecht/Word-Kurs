from __future__ import annotations

from pathlib import Path

from PIL import Image, ImageDraw
from docx import Document
from docx.shared import Pt, Cm

from course_common import (
    NAVY, TEAL, TEAL_DARK, PALE, PALE_TEAL, WARM, MID, TEXT, WHITE,
    base_doc, block, add_text, add_finish, finalise, set_run,
)
from course_build_helpers import clear_paragraph as _clear, style_run as _font


def practice_asset(path: Path):
    W, H = 1500, 900
    im = Image.new("RGB", (W, H), "#DCEBF2")
    d = ImageDraw.Draw(im)

    d.rectangle((0, 0, W, 470), fill="#DDECF4")
    d.polygon(
        [(0,470),(220,340),(430,430),(690,300),(930,430),(1190,325),(1500,455),(1500,610),(0,610)],
        fill="#7E9B8F",
    )
    d.rectangle((0, 560, W, H), fill="#AFCBB4")
    d.ellipse((120,470,1380,870), fill="#8FC4CB")
    d.ellipse((250,545,1260,810), fill="#A5D1D5")

    for x in [160,205,250,1260,1300,1340]:
        d.line((x,620,x-12,520), fill="#5F7D5D", width=5)
        d.line((x,620,x+8,535), fill="#5F7D5D", width=4)

    for x, y, scale in [(110,520,1.0),(1390,500,1.05)]:
        d.rectangle((x-13,y+35,x+13,y+205), fill="#775B45")
        d.ellipse((x-105*scale,y-90*scale,x+105*scale,y+100*scale), fill="#4F7566")
        d.ellipse((x-85*scale,y-145*scale,x+85*scale,y+40*scale), fill="#557E6D")

    d.polygon([(510,760),(950,760),(1040,900),(430,900)], fill="#D8C6A6")
    d.rectangle((630,690,860,715), fill="#8A6549")
    d.line((665,715,640,780), fill="#8A6549", width=12)
    d.line((825,715,850,780), fill="#8A6549", width=12)
    d.rectangle((4,4,W-5,H-5), outline="#D3DEE2", width=8)

    border = 140
    canvas = Image.new("RGB", (W + 2*border, H + 2*border), "white")
    canvas.paste(im, (border, border))
    canvas.save(path, quality=95)


def _add_req(cell, no, title, text, points):
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(.35)
    p.paragraph_format.line_spacing = 1
    r = p.add_run(f"{no}. ")
    set_run(r, size=9.1, bold=True, color=TEAL_DARK)
    r = p.add_run(title + " - ")
    set_run(r, size=9.1, bold=True, color=NAVY)
    r = p.add_run(text)
    set_run(r, size=9.1, color=TEXT)
    r = p.add_run(f"  [{points} P]")
    set_run(r, size=8.7, bold=True, color=MID)


def build_task_sheet(out: Path):
    doc = base_doc(
        "TEST",
        "Word",
        "So könnte der Test aussehen",
        "Richtzeit: 45 Minuten · 30 Punkte zur Orientierung · keine Note",
        "bekannte Word-Funktionen unter testähnlichen Bedingungen selbstständig anwenden.",
    )

    sec = doc.sections[0]
    ht = sec.header.tables[0]
    ht.cell(0,2).text = "ÜBUNGSTEST"
    for j, c in enumerate(ht.rows[0].cells):
        for p in c.paragraphs:
            for r in p.runs:
                _font(r, size=9.2, bold=True, color=WHITE if j == 0 else (TEAL_DARK if j == 2 else MID))

    ft = sec.footer.tables[0]
    ft.cell(0,1).text = "WORD · ÜBUNGSTEST"
    for c in ft.rows[0].cells:
        for p in c.paragraphs:
            for r in p.runs:
                _font(r, size=8.2, bold=True, color=MID)

    t = block(doc, "START", None, fill_left=NAVY, fill_right=PALE, label_color=WHITE, label_size=9.2)
    r = t.cell(0,1)
    p = r.paragraphs[0]
    _clear(p)
    x = p.add_run("Öffne ")
    set_run(x, size=9.6, bold=True, color=NAVY)
    x = p.add_run("Uebungstest_Ausgangsdokument.docx")
    set_run(x, size=9.6, bold=True, color=TEAL_DARK)
    x = p.add_run(" und speichere sofort eine Kopie als ")
    set_run(x, size=9.6, color=TEXT)
    x = p.add_run("Uebungstest_Nachname_Vorname.docx")
    set_run(x, size=9.6, bold=True, color=NAVY)
    add_text(r, "Arbeite nur in deiner Kopie. Die Bilddatei heisst uebungstest_greifensee.png.", 9.25, after=0)

    t = block(doc, "01-05", "GRUNDLAGEN")
    r = t.cell(0,1)
    _clear(r.paragraphs[0])
    _add_req(r, 1, "Seite", "A4 Hochformat und Seitenränder «Schmal».", 2)
    _add_req(r, 2, "Grundtext", "Ganzer Text Arial 11 pt; Zeilenabstand 1,15; Absatzabstand danach 6 pt.", 3)
    _add_req(r, 3, "Formatvorlagen", "Dokumenttitel = «Titel»; ÜBER DEN NATURTAG, MITBRINGEN, TAGESABLAUF, PROGRAMM und WICHTIG = «Überschrift 1».", 4)
    _add_req(r, 4, "Aufzählung", "Die fünf Dinge unter MITBRINGEN als echte Word-Aufzählung formatieren.", 2)
    _add_req(r, 5, "Nummerierung", "Die vier Schritte unter TAGESABLAUF als echte Word-Nummerierung formatieren.", 2)

    t = block(doc, "06-08", "BILD & TABELLE", fill_left=WARM, fill_right=WARM, label_color=NAVY)
    r = t.cell(0,1)
    _clear(r.paragraphs[0])
    _add_req(r, 6, "Bild", "Bilddatei einfügen; weissen Rand wegschneiden; Breite ca. 6 cm; Textumbruch «Quadrat»; rechts neben dem Einleitungstext platzieren.", 5)
    _add_req(r, 7, "Seitenumbruch", "Direkt vor PROGRAMM mit Ctrl + Enter eine neue Seite beginnen.", 2)
    _add_req(r, 8, "Tabelle", "Programmdaten in eine echte Tabelle mit 4 Spalten und 5 Zeilen übertragen; Kopfzeile fett und zentriert; Spalte «Tag» zentrieren.", 5)

    t = block(doc, "09-10", "ABSCHLUSS", fill_left=PALE_TEAL, fill_right=PALE_TEAL, label_color=TEAL_DARK)
    r = t.cell(0,1)
    _clear(r.paragraphs[0])
    _add_req(r, 9, "Kopfzeile", "Auf beiden Seiten oben «NATURTAG GREIFENSEE 2027» einfügen.", 2)
    _add_req(r, 10, "Fusszeile", "Unten mittig «Seite » plus automatische Seitenzahl einfügen. Nicht 1 und 2 von Hand tippen.", 3)

    t = block(doc, "CHECK", None, fill_left=WHITE, fill_right=WHITE, label_color=TEAL_DARK, label_size=9.2)
    r = t.cell(0,1)
    p = r.paragraphs[0]
    _clear(p)
    x = p.add_run("Vor dem Abgeben: ")
    set_run(x, size=9.4, bold=True, color=NAVY)
    x = p.add_run("Dateiname prüfen · genau 2 Seiten · Bild nicht verzerrt · echte Liste/Nummerierung/Tabelle · Kopfzeile auf beiden Seiten · Seitenzahlen automatisch.")
    set_run(x, size=9.15, color=TEXT)

    add_finish(doc, "Dieser Übungstest wird nicht benotet. Nutze Fehler danach gezielt zum Üben.")
    finalise(doc, out, "Word Übungstest")


def build_source_document(out: Path):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10.5)

    raw = [
        ("NATURTAG AM GREIFENSEE", 4),
        ("Donnerstag, 3. Juni 2027", 5),
        ("ÜBER DEN NATURTAG", 3),
        ("Unsere Klasse verbringt einen Tag am Greifensee. Wir beobachten Tiere und Pflanzen, arbeiten in kleinen Gruppen und halten unsere Ergebnisse in kurzen Notizen fest.", 4),
        ("Am Nachmittag vergleichen wir die Beobachtungen und machen einen gemeinsamen Rundgang am Ufer. Bei schlechtem Wetter wird das Programm angepasst.", 5),
        ("BILD HIER EINFÜGEN: uebungstest_greifensee.png", 5),
        ("MITBRINGEN", 3),
        ("Trinkflasche", 1),
        ("Lunch", 1),
        ("Regenjacke", 1),
        ("Schreibzeug", 1),
        ("kleines Sitzkissen", 5),
        ("TAGESABLAUF", 3),
        ("Treffpunkt beim Schulhaus", 1),
        ("Fahrt zum Greifensee", 1),
        ("Arbeit in Beobachtungsgruppen", 1),
        ("Gemeinsamer Abschluss und Rückfahrt", 5),
        ("PROGRAMM", 3),
        ("Tag | Zeit | Ort | Aktivität", 1),
        ("Donnerstag | 08.00 | Schulhaus | Treffpunkt", 1),
        ("Donnerstag | 09.15 | Greifensee | Beobachtungsauftrag", 1),
        ("Donnerstag | 12.00 | Uferplatz | Mittagspause", 1),
        ("Donnerstag | 15.15 | Bahnhof | Rückfahrt", 5),
        ("WICHTIG", 3),
        ("Wir bleiben während des Naturtags in unseren Gruppen. Abfälle nehmen wir wieder mit. Pflanzen werden nicht ausgerissen und Tiere nicht gestört.", 3),
    ]

    for text, after in raw:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(after)
        p.paragraph_format.line_spacing = 1
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(10.5)

    doc.core_properties.title = "Übungstest Ausgangsdokument"
    doc.core_properties.subject = "Word Kurs Sek 8"
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)


def build_all(root: Path):
    sheets = root / "arbeitsblaetter"
    assets = sheets / "assets"
    assets.mkdir(parents=True, exist_ok=True)

    practice_asset(assets / "uebungstest_greifensee.png")
    build_task_sheet(sheets / "Uebungstest_Word.docx")
    build_source_document(sheets / "Uebungstest_Ausgangsdokument.docx")