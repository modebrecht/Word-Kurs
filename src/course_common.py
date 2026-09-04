from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

NAVY = "17324D"
TEAL = "237B78"
TEAL_DARK = "1D6765"
PALE = "F3F6F7"
PALE_TEAL = "EAF4F3"
WARM = "F8F3EA"
MID = "667684"
LIGHT_LINE = "D3DEE2"
WHITE = "FFFFFF"
TEXT = "172A3A"

LEFT = Cm(3.2244)
RIGHT = Cm(14.1164)


def _font(run, name="Arial", size=None, bold=None, italic=None, color=None):
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.rFonts
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for a in ("ascii", "hAnsi", "eastAsia"):
        rFonts.set(qn(f"w:{a}"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _clear(p):
    for child in list(p._p):
        if child.tag != qn("w:pPr"):
            p._p.remove(child)


def _fill(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), color)


def _margins(cell, top=80, start=145, bottom=80, end=145):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for n, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        el = tcMar.find(qn(f"w:{n}"))
        if el is None:
            el = OxmlElement(f"w:{n}")
            tcMar.append(el)
        el.set(qn("w:w"), str(v))
        el.set(qn("w:type"), "dxa")


def _border(cell, **edges):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = tcPr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcPr.append(borders)
    for edge, attrs in edges.items():
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        for k, v in attrs.items():
            el.set(qn(f"w:{k}"), str(v))


def _keep_first(cell):
    while len(cell.paragraphs) > 1:
        p = cell.paragraphs[-1]
        p._element.getparent().remove(p._element)


def _grid(table):
    table.autofit = False
    table.columns[0].width = LEFT
    table.columns[1].width = RIGHT
    for row in table.rows:
        row.cells[0].width = LEFT
        row.cells[1].width = RIGHT
        for c in row.cells:
            c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _margins(row.cells[0], 75, 130, 75, 130)
        _margins(row.cells[1], 75, 190, 75, 190)


def _paragraph(p, text, size, bold=False, color=TEXT, after=0, align=None):
    _clear(p)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    _font(r, size=size, bold=bold, color=color)
    return p


def _label(cell, main, sub=None, fill=TEAL, color=WHITE, size=17):
    _keep_first(cell)
    _fill(cell, fill)
    p = cell.paragraphs[0]
    _clear(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(main)
    _font(r, size=size, bold=True, color=color)
    if sub:
        p.add_run("\n")
        r = p.add_run(sub)
        _font(r, size=8.1, bold=True, color=color)


def add_text(cell, text, size=9.6, bold=False, color=TEXT, after=0.5, italic=False):
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1
    r = p.add_run(text)
    _font(r, size=size, bold=bold, italic=italic, color=color)
    return p


def add_step(cell, letter, parts, after=0.45):
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1
    r = p.add_run(letter + "  ")
    _font(r, size=9.35, bold=True, color=TEAL)
    for item in parts:
        if isinstance(item, str):
            text, bold, italic, color = item, False, False, None
        else:
            text, bold, italic, color = item
        r = p.add_run(text)
        _font(r, size=9.35, bold=bold, italic=italic, color=color or TEXT)
    return p


def base_doc(code: str, topic: str, title: str, subtitle: str, learning_goal: str) -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(1.45)
    sec.bottom_margin = Cm(1.25)
    sec.left_margin = Cm(1.85)
    sec.right_margin = Cm(1.85)
    sec.header_distance = Cm(0.45)
    sec.footer_distance = Cm(0.55)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)

    h = sec.header.add_table(rows=1, cols=3, width=Cm(17.3))
    widths = [Cm(4.0), Cm(8.2), Cm(5.1)]
    for i, c in enumerate(h.rows[0].cells):
        c.width = widths[i]
        _margins(c, 70, 110, 70, 110)
    _fill(h.cell(0,0), NAVY)
    _fill(h.cell(0,1), PALE)
    _fill(h.cell(0,2), PALE_TEAL)
    texts = [("WORD KURS", WHITE), ("SEKUNDARSTUFE I · SEK 8", MID), (f"ARBEITSBLATT  {code}", TEAL_DARK)]
    for i,(txt,col) in enumerate(texts):
        p = h.cell(0,i).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(txt)
        _font(r, size=9.2, bold=True, color=col)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"{code}  ·  {topic.upper()}")
    _font(r, size=10, bold=True, color=TEAL_DARK)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    _font(r, size=27, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(subtitle)
    _font(r, size=10.8, color=MID)

    learning = block(doc, "LERNZIEL", None, fill_left=NAVY, fill_right=PALE, label_size=9.3)
    rcell = learning.cell(0,1)
    p = rcell.paragraphs[0]
    _clear(p)
    r = p.add_run("Ich kann ...")
    _font(r, size=12.6, bold=True, color=NAVY)
    add_text(rcell, learning_goal, 9.6, after=0)

    f = sec.footer.add_table(rows=1, cols=3, width=Cm(17.3))
    f.cell(0,0).text = "WORD KURS · SEK 8"
    f.cell(0,1).text = f"{code}  ·  {topic.upper()}"
    pnum = f.cell(0,2).paragraphs[0]
    _clear(pnum)
    pnum.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = pnum.add_run("SEITE ")
    _font(r, size=8.2, bold=True, color=MID)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    rnode = OxmlElement("w:r")
    tnode = OxmlElement("w:t")
    tnode.text = "1"
    rnode.append(tnode)
    fld.append(rnode)
    pnum._p.append(fld)
    for i,c in enumerate(f.rows[0].cells[:2]):
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                _font(r, size=8.2, bold=True, color=MID)
    return doc


def block(doc: Document, left: str, sub: str | None = None, *, fill_left=TEAL, fill_right=WHITE,
          label_color=WHITE, label_size=16.5):
    t = doc.add_table(rows=1, cols=2)
    _grid(t)
    _label(t.cell(0,0), left, sub, fill_left, label_color, label_size)
    _fill(t.cell(0,1), fill_right)
    return t


def add_spacer(doc, pts=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(pts)
    p.paragraph_format.line_spacing = 0.5
    return p


def add_workarea(doc, label_main, label_sub, lines, *, body_font="Times New Roman", size=10.2):
    t = block(doc, label_main, label_sub, fill_left=PALE, fill_right="FBFCFC", label_color=TEAL_DARK, label_size=11.0)
    l, r = t.rows[0].cells
    for c in (l,r):
        _border(c,
                top={"val":"single","sz":"5","color":LIGHT_LINE},
                bottom={"val":"single","sz":"5","color":LIGHT_LINE},
                left={"val":"single","sz":"5","color":LIGHT_LINE},
                right={"val":"single","sz":"5","color":LIGHT_LINE})
    _keep_first(r)
    p0 = r.paragraphs[0]
    _clear(p0)
    for i,line in enumerate(lines):
        p = p0 if i == 0 else r.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0.8)
        p.paragraph_format.line_spacing = 1
        rr = p.add_run(line)
        _font(rr, name=body_font, size=size, color="222222")
    return t


def add_tip(doc, text, lead="TIPP"):
    t = block(doc, lead, None, fill_left=WARM, fill_right=WARM, label_color=NAVY, label_size=9.3)
    p = t.cell(0,1).paragraphs[0]
    _clear(p)
    r = p.add_run(text)
    _font(r, size=9.5, color=TEXT)
    return t


def add_check(doc, text):
    t = block(doc, "CHECK", None, fill_left=PALE_TEAL, fill_right=PALE_TEAL, label_color=TEAL_DARK, label_size=9.1)
    p = t.cell(0,1).paragraphs[0]
    _clear(p)
    r = p.add_run(text)
    _font(r, size=9.45, color=TEXT)
    return t


def add_finish(doc, text='Gib dieses Arbeitsblatt in deinem Ordner "IB" ab.'):
    t = block(doc, "FERTIG?", None, fill_left=WHITE, fill_right=WHITE, label_color=TEAL_DARK, label_size=9.3)
    p = t.cell(0,1).paragraphs[0]
    _clear(p)
    r = p.add_run(text)
    _font(r, size=9.6, color=MID)
    for c in t.rows[0].cells:
        _border(c, top={"val":"single","sz":"7","color":LIGHT_LINE})
    return t


def add_picture(cell, path: Path, width_cm: float):
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0.5)
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    return p


def new_workspace_section(doc: Document, code: str, landscape=False):
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    if landscape:
        sec.orientation = WD_ORIENT.LANDSCAPE
        sec.page_width = Cm(29.7)
        sec.page_height = Cm(21)
    else:
        sec.page_width = Cm(21)
        sec.page_height = Cm(29.7)
    sec.top_margin = Cm(1.7)
    sec.bottom_margin = Cm(1.5)
    sec.left_margin = Cm(2.0)
    sec.right_margin = Cm(2.0)
    sec.header.is_linked_to_previous = True
    sec.footer.is_linked_to_previous = True
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(f"{code} · ÜBUNGSSEITE")
    _font(r, size=9.5, bold=True, color=TEAL_DARK)
    return sec


def finalise(doc: Document, path: Path, title: str):
    for table in doc.tables:
        for row in table.rows:
            trPr = row._tr.get_or_add_trPr()
            if trPr.find(qn("w:cantSplit")) is None:
                trPr.append(OxmlElement("w:cantSplit"))
    doc.core_properties.title = title
    doc.core_properties.subject = "Word Kurs Sek 8"
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def set_run(run, *, name="Arial", size=10.5, bold=False, italic=False, color=TEXT):
    _font(run, name=name, size=size, bold=bold, italic=italic, color=color)
