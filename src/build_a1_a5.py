from __future__ import annotations

from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from course_common import (
    NAVY, TEAL, TEAL_DARK, PALE, PALE_TEAL, WARM, MID, TEXT, WHITE,
    base_doc, block, add_text, add_step, add_workarea, add_tip, add_check,
    add_finish, add_picture, finalise, set_run, _clear, _keep_first, _fill,
)


def fonts():
    root = Path('/usr/share/fonts/truetype/liberation2')
    return str(root/'LiberationSans-Regular.ttf'), str(root/'LiberationSans-Bold.ttf')


def game_night_preview(path: Path):
    reg, bold = fonts()
    W,H=1500,540
    im=Image.new('RGB',(W,H),'white'); d=ImageDraw.Draw(im)
    f_title=ImageFont.truetype(bold,66); f_sub=ImageFont.truetype(bold,34)
    f_body=ImageFont.truetype(reg,30); f_body_b=ImageFont.truetype(bold,30); f_small=ImageFont.truetype(reg,28)
    d.rounded_rectangle((18,18,W-18,H-18),radius=18,outline='#D3DEE2',width=3,fill='white')
    x=65; y=52
    d.text((x,y),'GAME NIGHT',font=f_title,fill='#237B78'); y+=88
    d.line((x,y,W-65,y),fill='#237B78',width=5); y+=30
    d.text((x,y),'Freitag, 18. September',font=f_sub,fill='#17324D'); y+=58
    d.text((x,y),'18.30-21.00 Uhr',font=f_body_b,fill='#17324D')
    d.text((515,y),'Zimmer 204',font=f_body_b,fill='#237B78'); y+=64
    d.text((x,y),'Bring dein Lieblingsspiel mit.',font=f_body,fill='#17324D'); y+=52
    d.text((x,y),'Getränke stehen bereit. Anmeldung bis Mittwoch bei Frau Keller.',font=f_small,fill='#5E6D78')
    path.parent.mkdir(parents=True, exist_ok=True); im.save(path,quality=95)


def kino_preview(path: Path):
    reg,bold=fonts(); W,H=1500,470
    im=Image.new('RGB',(W,H),'white'); d=ImageDraw.Draw(im)
    ft=ImageFont.truetype(bold,54); fd=ImageFont.truetype(bold,28); fb=ImageFont.truetype(reg,27); fbb=ImageFont.truetype(bold,27)
    d.rounded_rectangle((18,18,W-18,H-18),radius=18,outline='#D3DEE2',width=3,fill='white')
    s='KINOABEND'; box=d.textbbox((0,0),s,font=ft); d.text(((W-(box[2]-box[0]))/2,42),s,font=ft,fill='#237B78')
    s='Freitag, 25. September'; box=d.textbbox((0,0),s,font=fd); d.text((W-70-(box[2]-box[0]),120),s,font=fd,fill='#17324D')
    y=205
    for s,f,c in [('Wir schauen gemeinsam einen Film in der Aula.',fb,'#17324D'),('Treffpunkt: 18.15 Uhr vor der Aula.',fbb,'#17324D'),('Ende: ca. 21.00 Uhr.',fbb,'#17324D'),('Bring etwas zu trinken mit.',fb,'#5E6D78')]:
        d.text((70,y),s,font=f,fill=c); y+=58
    im.save(path,quality=95)


def list_preview(path: Path):
    reg,bold=fonts(); W,H=1500,390
    im=Image.new('RGB',(W,H),'white'); d=ImageDraw.Draw(im)
    ft=ImageFont.truetype(bold,48); fh=ImageFont.truetype(bold,28); fb=ImageFont.truetype(reg,25)
    d.rounded_rectangle((18,18,W-18,H-18),radius=18,outline='#D3DEE2',width=3,fill='white')
    d.text((60,36),'KLASSENAUSFLUG',font=ft,fill='#237B78')
    d.text((60,112),'Mitnehmen',font=fh,fill='#17324D'); y=155
    for s in ['Trinkflasche','Lunch','Regenjacke']:
        d.text((68,y),'•',font=fb,fill='#237B78'); d.text((105,y),s,font=fb,fill='#17324D'); y+=42
    x=760; d.text((x,112),'So läuft es ab',font=fh,fill='#17324D'); y=155
    for i,s in enumerate(['Treffpunkt beim Schulhaus','Zugfahrt','Wanderung zum Aussichtspunkt'],1):
        d.text((x+4,y),f'{i}.',font=fb,fill='#237B78'); d.text((x+45,y),s,font=fb,fill='#17324D'); y+=42
    im.save(path,quality=95)


def school_icon(path: Path):
    W,H=420,300
    im=Image.new('RGBA',(W,H),(255,255,255,0)); d=ImageDraw.Draw(im)
    d.rounded_rectangle((55,95,365,260),radius=12,fill='#EAF4F3',outline='#237B78',width=5)
    d.polygon([(45,105),(210,35),(375,105)],fill='#237B78')
    d.rounded_rectangle((180,170,240,260),radius=6,fill='#17324D')
    for x in (95,275): d.rounded_rectangle((x,145,x+55,195),radius=5,fill='white',outline='#17324D',width=4)
    d.line((210,35,210,2),fill='#17324D',width=5); d.polygon([(210,4),(274,20),(210,36)],fill='#17324D')
    im.save(path)


def chaos_preview(path: Path, icon_path: Path):
    reg,bold=fonts(); W,H=1500,520
    im=Image.new('RGB',(W,H),'white'); d=ImageDraw.Draw(im)
    ft=ImageFont.truetype(bold,50); fs=ImageFont.truetype(bold,27); fb=ImageFont.truetype(reg,23); fh=ImageFont.truetype(bold,24)
    d.rounded_rectangle((18,18,W-18,H-18),radius=18,outline='#D3DEE2',width=3,fill='white')
    icon=Image.open(icon_path).convert('RGBA'); icon.thumbnail((115,82)); im.paste(icon,((W-icon.width)//2,34),icon)
    s='SCHULFEST 2026'; box=d.textbbox((0,0),s,font=ft); d.text(((W-(box[2]-box[0]))/2,118),s,font=ft,fill='#17324D')
    s='Freitag, 12. Juni · 17.30 Uhr'; box=d.textbbox((0,0),s,font=fs); d.text(((W-(box[2]-box[0]))/2,188),s,font=fs,fill='#237B78')
    d.text((80,255),'Mitnehmen',font=fh,fill='#17324D'); y=300
    for s in ['Trinkflasche','Jacke für den Abend','gute Laune']:
        d.text((88,y),'•',font=fb,fill='#237B78'); d.text((125,y),s,font=fb,fill='#17324D'); y+=43
    x=760; d.text((x,255),'Ablauf',font=fh,fill='#17324D'); y=300
    for i,s in enumerate(['Begrüssung','Spiel & Essen','Musik in der Aula'],1):
        d.text((x+5,y),f'{i}.',font=fb,fill='#237B78'); d.text((x+48,y),s,font=fb,fill='#17324D'); y+=43
    im.save(path,quality=95)


def build_a1(out: Path):
    doc=base_doc('A1','Text formatieren','Die wichtigsten Werkzeuge','Du lernst die Grundwerkzeuge zum Formatieren von Text.','Text markieren und Schriftart, Schriftgrösse, Fett, Kursiv und Farbe einsetzen und eine Änderung rückgängig machen.')
    t=block(doc,'01','AUFGABE'); r=t.cell(0,1); p=r.paragraphs[0]; _clear(p); x=p.add_run('Formatiere Schritt für Schritt'); set_run(x,size=12.8,bold=True,color=NAVY)
    add_text(r,'Arbeite von A bis E. Markiere immer nur den Text, den der Schritt nennt.',9.4,after=.7)
    p=r.add_paragraph(); x=p.add_run('So sieht es aus:  '); set_run(x,size=9.3,bold=True,color=MID); x=p.add_run('Fett'); set_run(x,size=9.3,bold=True,color=NAVY); x=p.add_run('   ·   '); set_run(x,size=9.3,color=MID); x=p.add_run('Kursiv'); set_run(x,size=9.3,italic=True,color=NAVY); x=p.add_run('   ·   '); set_run(x,size=9.3,color=MID); x=p.add_run('Blau'); set_run(x,size=9.3,bold=True,color=TEAL)
    add_step(r,'A',[( 'Markiere zuerst den ganzen Rohtext',True,False,NAVY),('  →  Arial  →  11 pt',False,False,None)])
    add_step(r,'B',[('Markiere nur «Mein Lieblingsort»',True,False,NAVY),('  →  20 pt  →  ',False,False,None),('Fett',True,False,NAVY),('  →  Dunkelblau',False,False,NAVY)])
    add_step(r,'C',[('Markiere «Walensee» und danach «Churfirsten»',True,False,NAVY),('  →  ',False,False,None),('Fett',True,False,NAVY)])
    add_step(r,'D',[('Markiere nur «Sommer»',True,False,NAVY),('  →  ',False,False,None),('Blau',True,False,TEAL)])
    add_step(r,'E',[('Markiere den letzten Satz',True,False,NAVY),('  →  ',False,False,None),('Kursiv',False,True,NAVY)])
    add_workarea(doc,'HIER','ARBEITEN 01',['Mein Lieblingsort','Der Walensee liegt in der Ostschweiz. Auf der einen Seite ragen die Churfirsten steil auf, auf der anderen liegt das Glarnerland.','Im Sommer kann man baden, wandern oder einfach am Ufer sitzen.','Besonders schön finde ich den Blick über den See am Abend.'])
    add_tip(doc,'Etwas falsch gemacht? Mit Ctrl + Z machst du den letzten Schritt rückgängig.')
    add_check(doc,'Sieht dein Titel deutlich anders aus als der Fliesstext? Sind «Walensee» und «Churfirsten» fett?')
    add_finish(doc); finalise(doc,out,'A1 - Text formatieren')


def build_a2(out: Path, preview: Path):
    doc=base_doc('A2','Nach Vorlage gestalten','Schau genau hin','Du nutzt die Werkzeuge aus A1 und baust eine sichtbare Vorlage nach.','an einer Vorlage erkennen, welche Textstellen grösser, fett oder farbig formatiert sind, und diese Formatierung nachbauen.')
    t=block(doc,'01','NACHBAUEN'); r=t.cell(0,1); p=r.paragraphs[0]; _clear(p); x=p.add_run('Baue die Vorlage genau nach'); set_run(x,size=12.8,bold=True,color=NAVY)
    add_text(r,'Verändere den Text nicht. Verändere nur die Formatierung.',9.5,after=.6); add_picture(r,preview,11.7)
    add_workarea(doc,'HIER','ARBEITEN 01',['Game Night','Freitag, 18. September','18.30-21.00 Uhr','Zimmer 204','Bring dein Lieblingsspiel mit.','Getränke stehen bereit. Anmeldung bis Mittwoch bei Frau Keller.'])
    add_tip(doc,'Nicht raten – vergleichen. Schau immer nur auf eine Sache: zuerst Grösse, dann Fett, dann Farbe.')
    add_check(doc,'Vergleiche Zeile für Zeile mit der Vorlage. Findest du noch einen Unterschied?')
    add_finish(doc); finalise(doc,out,'A2 - Nach Vorlage gestalten')


def build_a3(out: Path, preview: Path):
    doc=base_doc('A3','Absätze & Ordnung','Text braucht Luft','Du lernst, wie du Absätze ausrichtest und Abstände sauber einstellst.','Absätze links, zentriert oder rechts ausrichten und Zeilen- sowie Absatzabstände einstellen.')
    t=block(doc,'01','AUFGABE'); r=t.cell(0,1); p=r.paragraphs[0]; _clear(p); x=p.add_run('Ordne den Infotext Schritt für Schritt'); set_run(x,size=12.8,bold=True,color=NAVY)
    add_text(r,'Markiere immer genau den Absatz oder die Absätze, die im Schritt genannt werden.',9.4,after=.6)
    add_step(r,'A',[('«Bibliothek am Mittag»',True,False,NAVY),('  →  zentriert',False,False,None)])
    add_step(r,'B',[('«Schulhaus Sonnenberg»',True,False,NAVY),('  →  rechts',False,False,None)])
    add_step(r,'C',[('Die drei Textabsätze',True,False,NAVY),('  →  links',False,False,None)])
    add_step(r,'D',[('Die drei Textabsätze',True,False,NAVY),('  →  Zeilenabstand 1,5',False,False,None)])
    add_step(r,'E',[('Die drei Textabsätze',True,False,NAVY),('  →  Abstand nach Absatz: 6 pt',False,False,None)])
    add_workarea(doc,'HIER','ARBEITEN 01',['Bibliothek am Mittag','Schulhaus Sonnenberg','Unsere Bibliothek ist am Dienstag und Donnerstag über Mittag geöffnet.','Du kannst lesen, Hausaufgaben erledigen oder in Ruhe arbeiten.','Bitte stelle Bücher nach dem Lesen zurück und verlasse den Raum ordentlich.'])
    add_tip(doc,'Für Abstand nicht mehrfach Enter drücken. Stelle den Abstand beim Absatz ein. So bleibt das Dokument sauber.')
    t=block(doc,'02','NACHBAUEN',fill_left=TEAL_DARK); r=t.cell(0,1); p=r.paragraphs[0]; _clear(p); x=p.add_run('Baue die Vorlage nach'); set_run(x,size=12.8,bold=True,color=NAVY); add_text(r,'Verändere den Text nicht. Nutze nur Ausrichtung und Abstände.',9.4); add_picture(r,preview,11.1)
    add_check(doc,'Hast du Leerzeilen eingefügt? Falls ja: löschen. Die Abstände sollen durch Absatz-Einstellungen entstehen.')
    add_finish(doc); finalise(doc,out,'A3 - Absätze und Ordnung')


def build_a4(out: Path, preview: Path):
    doc=base_doc('A4','Listen & Nummerierungen','Dinge oder Reihenfolge?','Du lernst, wann du Aufzählungen und wann du Nummerierungen verwendest.','Dinge als Aufzählung und eine Reihenfolge als Nummerierung darstellen.')
    t=block(doc,'01','AUFGABE'); r=t.cell(0,1); p=r.paragraphs[0]; _clear(p); x=p.add_run('Mach aus Zeilen richtige Listen'); set_run(x,size=12.8,bold=True,color=NAVY)
    add_text(r,'Dinge = Aufzählung. Reihenfolge = Nummerierung.',9.5,bold=True,color=MID)
    add_step(r,'A',[('Markiere die drei Dinge unter «Mitnehmen»',True,False,NAVY),('  →  Aufzählung',False,False,None)])
    add_step(r,'B',[('Markiere die drei Schritte unter «Ablauf»',True,False,NAVY),('  →  Nummerierung',False,False,None)])
    add_step(r,'C',[('Füge unter «Mitnehmen» «Sonnencreme» hinzu',True,False,NAVY),('  →  sie gehört zur Aufzählung',False,False,None)])
    add_workarea(doc,'HIER','ARBEITEN 01',['SPORTTAG','Mitnehmen','Turnschuhe','Trinkflasche','Lunch','Ablauf','Treffpunkt beim Schulhaus','Gemeinsames Aufwärmen','Start der Wettkämpfe'],size=9.8)
    add_tip(doc,'Punkte und Zahlen nicht von Hand tippen. Benutze die Word-Funktion für Aufzählung oder Nummerierung.')
    t=block(doc,'02','NACHBAUEN',fill_left=TEAL_DARK); r=t.cell(0,1); p=r.paragraphs[0]; _clear(p); x=p.add_run('Baue die Listen der Vorlage nach'); set_run(x,size=12.8,bold=True,color=NAVY); add_text(r,'Verändere den Text nicht. Entscheide anhand der Vorlage: Punkte oder Zahlen?',9.4); add_picture(r,preview,10.5)
    add_check(doc,'Lösche in der Nummerierung den zweiten Schritt. Prüfe, ob Word automatisch neu nummeriert.')
    add_finish(doc); finalise(doc,out,'A4 - Listen und Nummerierungen')


def build_a5(out: Path, preview: Path, icon: Path):
    doc=base_doc('A5','Dokument retten','Rette das Chaos-Dokument','Vergleiche mit der Vorlage und repariere gezielte Formatierungsfehler.','bekannte Formatierungsfehler erkennen und reparieren. Eine vorhandene Grafik kann ich verkleinern und an die richtige Stelle verschieben.')
    t=block(doc,'01','REPARIEREN'); r=t.cell(0,1); p=r.paragraphs[0]; _clear(p); x=p.add_run('Mach das Chaos wieder ordentlich'); set_run(x,size=12.8,bold=True,color=NAVY); add_text(r,'Vergleiche immer wieder mit der Vorlage. Arbeite von A bis F.',9.2); add_picture(r,preview,9.7)
    add_step(r,'A',[('Titel «SCHULFEST 2026»',True,False,NAVY),('  →  Arial, 20 pt, fett, dunkelblau, zentriert',False,False,None)])
    add_step(r,'B',[('Datum',True,False,NAVY),('  →  Arial, 11 pt, fett, zentriert',False,False,None)])
    add_step(r,'C',[('«Mitnehmen» + drei Dinge',True,False,NAVY),('  →  Überschrift fett, Dinge als Aufzählung',False,False,None)])
    add_step(r,'D',[('«Ablauf» + drei Schritte',True,False,NAVY),('  →  Überschrift fett, Schritte nummerieren',False,False,None)])
    add_step(r,'E',[('Normale Textzeilen',True,False,NAVY),('  →  Arial, 11 pt, links, Zeilenabstand 1,0',False,False,None)])
    add_step(r,'F',[('Grafik',True,False,NAVY),('  →  ca. 2,2 cm breit, direkt über dem Titel, zentriert',False,False,None)])
    t=block(doc,'CHAOS','DOKUMENT 01',fill_left=PALE,fill_right='FBFCFC',label_color=TEAL_DARK,label_size=11.0); r=t.cell(0,1); p=r.paragraphs[0]; _clear(p)
    x=p.add_run('SCHULFEST 2026'); set_run(x,name='Times New Roman',size=24,italic=True,color='C05A2B')
    p=r.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.RIGHT; x=p.add_run('Freitag, 12. Juni · 17.30 Uhr'); set_run(x,name='Courier New',size=9,color='7A2B83')
    p=r.add_paragraph(); p.paragraph_format.space_after=Pt(0); p.add_run().add_picture(str(icon),width=Cm(3.3))
    for text,name,size,bold,italic,color,align in [
        ('Mitnehmen','Arial',13,False,True,'C05A2B',WD_ALIGN_PARAGRAPH.RIGHT),
        ('1. Trinkflasche','Comic Sans MS',11,False,False,'333333',WD_ALIGN_PARAGRAPH.LEFT),
        ('2. Jacke für den Abend','Comic Sans MS',11,False,False,'333333',WD_ALIGN_PARAGRAPH.LEFT),
        ('3. gute Laune','Comic Sans MS',11,False,False,'333333',WD_ALIGN_PARAGRAPH.LEFT),
        ('Ablauf','Arial',13,False,False,TEAL,WD_ALIGN_PARAGRAPH.CENTER),
        ('• Begrüssung','Times New Roman',11,False,False,'333333',WD_ALIGN_PARAGRAPH.LEFT),
        ('• Spiel & Essen','Times New Roman',11,False,False,'333333',WD_ALIGN_PARAGRAPH.LEFT),
        ('• Musik in der Aula','Times New Roman',11,False,False,'333333',WD_ALIGN_PARAGRAPH.LEFT),
    ]:
        p=r.add_paragraph(); p.alignment=align; p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=.9; x=p.add_run(text); set_run(x,name=name,size=(10 if size>=11 else size),bold=bold,italic=italic,color=color)
    add_tip(doc,'Grafik verschieben: anklicken → Ctrl + X → Cursor über den Titel → Ctrl + V. Danach zentrieren und verkleinern.  FERTIG? Gib das Blatt in deinem Ordner "IB" ab.')
    finalise(doc,out,'A5 - Rette das Chaos-Dokument')


def build_all(root: Path):
    sheets=root/'arbeitsblaetter'; assets=sheets/'assets'; prev=assets/'vorlagen'; assets.mkdir(parents=True,exist_ok=True); prev.mkdir(parents=True,exist_ok=True)
    game=prev/'a2_game_night_vorlage.png'; kino=prev/'a3_kinoabend_vorlage.png'; lists=prev/'a4_ausflug_vorlage.png'; icon=assets/'a5_school_icon.png'; chaos=prev/'a5_schulfest_vorlage.png'
    game_night_preview(game); kino_preview(kino); list_preview(lists); school_icon(icon); chaos_preview(chaos,icon)
    build_a1(sheets/'A1_Text_formatieren.docx')
    build_a2(sheets/'A2_Nach_Vorlage_gestalten.docx',game)
    build_a3(sheets/'A3_Absaetze_und_Ordnung.docx',kino)
    build_a4(sheets/'A4_Listen_und_Nummerierungen.docx',lists)
    build_a5(sheets/'A5_Rette_das_Chaos_Dokument.docx',chaos,icon)
