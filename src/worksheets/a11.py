from __future__ import annotations

import sys
from pathlib import Path

if __package__ in (None, ""):
    sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from PIL import Image, ImageDraw, ImageFont
from docx.shared import Pt

from build_runtime import resolve_font_paths
from course_common import (
    NAVY, TEAL_DARK, WARM,
    base_doc, block, add_text, add_tip, add_check, add_finish,
    add_picture, finalise, set_run,
)
from course_build_helpers import clear_paragraph as _clear, new_detached_workspace_section


def asset(path: Path):
    W,H=1200,800
    im=Image.new('RGB',(W,H),'#F7FBFC'); d=ImageDraw.Draw(im)
    d.rectangle((0,0,W,220),fill='#DCECF5'); d.rectangle((0,220,W,520),fill='#EAF4F8'); d.rectangle((0,520,W,H),fill='#DCE9E3')
    d.ellipse((110,540,1090,770),fill='#A9D2D5'); d.ellipse((180,585,1020,720),fill='#96C3C7')
    d.polygon([(70,565),(370,185),(675,565)],fill='#6E828E'); d.polygon([(390,565),(735,115),(1100,565)],fill='#566C78')
    d.polygon([(245,355),(370,185),(500,355)],fill='white'); d.polygon([(585,305),(735,115),(900,315)],fill='white')
    d.line([(70,565),(370,185),(675,565)],fill='#4F6672',width=4); d.line([(390,565),(735,115),(1100,565)],fill='#415863',width=4)
    for x,y,scale in [(125,495,1.0),(1015,485,1.08),(935,525,.82),(235,535,.76)]:
        d.rectangle((x-8,y+35,x+8,y+102),fill='#7A5B43')
        d.polygon([(x,y-50*scale),(x-55*scale,y+52*scale),(x+55*scale,y+52*scale)],fill='#4A7568')
        d.polygon([(x,y-5*scale),(x-48*scale,y+75*scale),(x+48*scale,y+75*scale)],fill='#3D6257')
    d.rectangle((4,4,W-5,H-5),outline='#D3DEE2',width=8); im.save(path,quality=95)


def preview(path: Path, asset_path: Path):
    reg,bold=resolve_font_paths(); W,H=1600,1040
    im=Image.new('RGB',(W,H),'white'); d=ImageDraw.Draw(im)
    fh=ImageFont.truetype(bold,28); ft=ImageFont.truetype(bold,58); f1=ImageFont.truetype(bold,36); fb=ImageFont.truetype(reg,31); fbb=ImageFont.truetype(bold,30); fs=ImageFont.truetype(reg,26); ftab=ImageFont.truetype(reg,27); ftabb=ImageFont.truetype(bold,27)
    d.rounded_rectangle((18,18,W-18,H-18),radius=18,outline='#D3DEE2',width=3,fill='white')
    d.text((65,45),'KLASSENLAGER 2026',font=fh,fill='#17324D'); d.text((1300,45),'SEK 8B',font=fh,fill='#237B78')
    d.text((65,125),'KLASSENLAGER FLIMS',font=ft,fill='#17324D'); d.text((67,200),'15.–18. Juni 2026',font=fbb,fill='#237B78')
    pic=Image.open(asset_path).convert('RGB'); pic.thumbnail((520,350)); im.paste(pic,(1010,145)); d.rectangle((1010,145,1010+pic.width,145+pic.height),outline='#D3DEE2',width=2)
    d.text((65,285),'Vier Tage unterwegs in Graubünden:',font=fb,fill='#17324D'); d.text((65,335),'wandern, gemeinsam kochen und Zeit am See.',font=fb,fill='#17324D')
    d.text((65,425),'MITNEHMEN',font=f1,fill='#237B78'); y=480
    for item in ['Wanderschuhe','Regenjacke','Trinkflasche','kleiner Rucksack']:
        d.ellipse((74,y+10,88,y+24),fill='#17324D'); d.text((105,y),item,font=fb,fill='#17324D'); y+=50
    d.text((65,700),'PROGRAMM',font=f1,fill='#237B78')
    x0,y0=65,755; cols=[220,390,420]; rows=[['Tag','Vormittag','Nachmittag'],['Montag','Anreise','Dorfrundgang'],['Dienstag','Wanderung','Caumasee'],['Mittwoch','Sport','Freizeit'],['Donnerstag','Aufräumen','Rückreise']]
    y=y0
    for ri,row in enumerate(rows):
        x=x0
        row_h=52 if ri==0 else 49
        for w,val in zip(cols,row):
            d.rectangle((x,y,x+w,y+row_h),fill='white',outline='#D3DEE2',width=2); d.text((x+12,y+12),val,font=ftabb if ri==0 else ftab,fill='#17324D'); x+=w
        y+=row_h
    d.text((65,1000),'Klassenlager Flims',font=fs,fill='#667684'); d.text((1450,1000),'1',font=fs,fill='#667684')
    im.save(path,quality=95)


def build_document(out: Path, preview_path: Path):
    doc=base_doc('A11','Dokument nachbauen','Jetzt ohne Klickanleitung','Du kombinierst bekannte Word-Werkzeuge und baust eine Vorlage möglichst genau nach.','mehrere bekannte Word-Funktionen selbstständig kombinieren und ein Dokument nach einer sichtbaren Vorlage nachbauen.')
    t=block(doc,'01','NACHBAUEN'); r=t.cell(0,1); p=r.paragraphs[0]; _clear(p); x=p.add_run('Baue Seite 2 so um, dass sie wie diese Vorlage aussieht'); set_run(x,size=12.6,bold=True,color=NAVY); add_text(r,'Der Text und die Daten sind vorgegeben. Du entscheidest selbst, welche bekannten Werkzeuge du dafür brauchst.',9.35); add_picture(r,preview_path,10.9)
    t=block(doc,'PFLICHT',None,fill_left=WARM,fill_right=WARM,label_color=NAVY,label_size=9.3); p=t.cell(0,1).paragraphs[0]; _clear(p); x=p.add_run('Verwende wirklich: '); set_run(x,size=9.55,bold=True,color=NAVY); x=p.add_run('Formatvorlagen für Titel/Überschriften · echte Aufzählung · echte Tabelle · Bild mit Textumbruch · Kopfzeile · automatische Seitenzahl.'); set_run(x,size=9.25)
    add_tip(doc,'Bei Titel und Überschriften zählt die richtige Formatvorlage. Das genaue Aussehen darf je nach Word-Version leicht von der Vorlage abweichen.','MERKE')
    add_tip(doc,'Arbeite von gross nach klein: zuerst Aufbau und Bereiche, danach Bild/Tabelle, ganz am Schluss Abstände und Feinarbeit.')
    add_check(doc,'Stimmen Reihenfolge, Grössenverhältnisse, Bildposition, Liste, Tabelle sowie Kopf- und Fussbereich?')
    add_finish(doc,'Gib dieses Arbeitsblatt zusammen mit der Bilddatei in deinem Ordner "IB" ab.')
    new_detached_workspace_section(doc,'A11',top_margin_cm=1.8,bottom_margin_cm=1.8)
    for text in ['KLASSENLAGER FLIMS','15.–18. Juni 2026','Vier Tage unterwegs in Graubünden: wandern, gemeinsam kochen und Zeit am See.','MITNEHMEN','Wanderschuhe','Regenjacke','Trinkflasche','kleiner Rucksack','PROGRAMM','Tag | Vormittag | Nachmittag','Montag | Anreise | Dorfrundgang','Dienstag | Wanderung | Caumasee','Mittwoch | Sport | Freizeit','Donnerstag | Aufräumen | Rückreise']:
        p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(3.5); x=p.add_run(text); set_run(x,size=11)
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(8); x=p.add_run('BILDDATEI: a11_klassenlager_berge.png'); set_run(x,size=9.5,bold=True,color=TEAL_DARK)
    finalise(doc,out,'A11 - Dokument nach Vorlage nachbauen')


def build(root: Path):
    sheets=root/'arbeitsblaetter'; assets=sheets/'assets'; prev=assets/'vorlagen'; assets.mkdir(parents=True,exist_ok=True); prev.mkdir(parents=True,exist_ok=True)
    image_asset=assets/'a11_klassenlager_berge.png'; image_preview=prev/'a11_klassenlager_vorlage.png'
    asset(image_asset); preview(image_preview,image_asset)
    build_document(sheets/'A11_Dokument_nach_Vorlage_nachbauen.docx',image_preview)


if __name__ == '__main__':
    from worksheet_runtime import run_single
    run_single('A11', build)
