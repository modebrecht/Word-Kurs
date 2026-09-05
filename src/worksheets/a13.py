from __future__ import annotations

import sys
from pathlib import Path

if __package__ in (None, ""):
    sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from PIL import Image, ImageDraw
from docx.shared import Pt

from course_common import (
    NAVY, TEAL_DARK, PALE_TEAL, WARM, TEXT,
    base_doc, block, add_text, add_check, add_finish, finalise, set_run,
)
from course_build_helpers import clear_paragraph as _clear, new_detached_workspace_section


def asset(path: Path):
    SW, SH = 1600, 950
    scene = Image.new("RGB", (SW, SH), "#E7F1F5")
    d = ImageDraw.Draw(scene)
    d.rectangle((0, 0, SW, 560), fill="#DDECF3")
    d.polygon(
        [(0,560),(280,400),(520,510),(760,360),(1030,500),(1320,350),(1600,520),(1600,700),(0,700)],
        fill="#8DA59A",
    )
    d.polygon([(0,650),(360,610),(820,665),(1220,625),(1600,690),(1600,950),(0,950)], fill="#9DC9D0")
    buildings = [
        (110,430,270,680,"#D7B18A"),(285,400,450,680,"#E3C39E"),(465,455,635,680,"#D9AA84"),
        (650,390,825,680,"#E5C8A8"),(850,445,1030,680,"#C99D7B"),(1045,410,1215,680,"#E1BE94"),(1230,450,1430,680,"#D4A77F"),
    ]
    for x1, y1, x2, y2, color in buildings:
        d.rectangle((x1,y1,x2,y2), fill=color, outline="#17324D", width=4)
        d.polygon([(x1-8,y1),(x2+8,y1),((x1+x2)//2,y1-75)], fill="#7A5F55", outline="#17324D")
        for yy in range(y1+55, y2-40, 85):
            for xx in range(x1+35, x2-25, 60):
                d.rectangle((xx,yy,xx+28,yy+38), fill="#EAF4F3", outline="#17324D", width=2)
    tx1, tx2 = 720, 875
    d.rectangle((tx1,230,tx2,680), fill="#C6AA7D", outline="#17324D", width=5)
    d.polygon([(700,230),(895,230),(798,125)], fill="#4D6D67", outline="#17324D")
    d.ellipse((752,305,843,396), fill="#F4F0E7", outline="#17324D", width=4)
    d.line((797,350,797,320), fill="#17324D", width=4); d.line((797,350,825,365), fill="#17324D", width=4)
    d.polygon([(40,690),(680,735),(680,780),(40,735)], fill="#8A7B6C", outline="#17324D")
    for x in range(95,650,90): d.line((x,720,x,760), fill="#E8E1D4", width=9)
    for x in (75,1510):
        d.rectangle((x-12,570,x+12,755), fill="#775B45"); d.ellipse((x-95,470,x+95,645), fill="#4F7667")
    d.rectangle((4,4,SW-5,SH-5), outline="#D3DEE2", width=8)

    # Deliberately give the source image generous empty margins so that
    # cropping remains a meaningful transfer task in the final exercise.
    W, H = 1960, 1250
    im = Image.new("RGB", (W, H), "#F4F6F7")
    im.paste(scene, (180, 150))
    im.save(path, quality=95)


def _requirement(cell, title, text):
    p = cell.add_paragraph(); p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(.55); p.paragraph_format.line_spacing = 1
    r = p.add_run(title + ": "); set_run(r, size=9.25, bold=True, color=NAVY)
    r = p.add_run(text); set_run(r, size=9.25, color=TEXT)


def build_document(out: Path):
    doc = base_doc(
        "A13", "Gesamtauftrag", "Alles zusammen",
        "Keine neue Word-Funktion mehr. Jetzt kombinierst du das Gelernte selbstständig.",
        "bekannte Word-Werkzeuge selbstständig auswählen und daraus ein sauberes, zweiseitiges Dokument erstellen.",
    )
    t = block(doc, "01", "GESAMTAUFTRAG"); r = t.cell(0,1); p = r.paragraphs[0]; _clear(p)
    x = p.add_run("Gestalte ab Seite 2 ein Infodokument «PROJEKTWOCHE BERN»"); set_run(x, size=12.5, bold=True, color=NAVY)
    add_text(r, "Der Inhalt ist vollständig vorgegeben. Dein fertiges Infodokument soll zwei Seiten lang sein. Falls deine Word-Version bei Formatvorlagen leicht andere Abstände verwendet, darfst du kleine Abstände anpassen – Schrift und Inhalte nicht verkleinern.", 9.35)
    t = block(doc, "PFLICHT", "AUFBAU", fill_left=WARM, fill_right=WARM, label_color=NAVY, label_size=12.7); r = t.cell(0,1); _clear(r.paragraphs[0])
    _requirement(r, "Seite", "A4 Hochformat · Seitenränder «Schmal»")
    _requirement(r, "Fliesstext", "Arial 11 pt · Zeilenabstand 1,15 · Absatzabstand danach 6 pt")
    _requirement(r, "Struktur", "Titel mit Formatvorlage «Titel» · Hauptbereiche mit «Überschrift 1»")
    _requirement(r, "Umbruch", "Vor «PROGRAMM» einen echten Seitenumbruch mit Ctrl + Enter setzen")
    t = block(doc, "PFLICHT", "ELEMENTE", fill_left=PALE_TEAL, fill_right=PALE_TEAL, label_color=TEAL_DARK, label_size=12.7); r = t.cell(0,1); _clear(r.paragraphs[0])
    _requirement(r, "Bild", "a13_bern_altstadt.png einfügen · zuschneiden · ca. 6 cm breit · Textumbruch «Quadrat»")
    _requirement(r, "Listen", "«MITBRINGEN» als echte Aufzählung · «TAGESABLAUF» als echte Nummerierung")
    _requirement(r, "Tabelle", "Programmdaten als echte Tabelle mit 4 Spalten darstellen · Kopfzeile fett")
    _requirement(r, "Kopf/Fuss", "Kopfzeile «PROJEKTWOCHE BERN» · Fusszeile mit automatischer Seitenzahl")
    add_check(doc,"Kontrolliere jede Pflichtzeile einzeln. Wenn alle erfüllt sind und das Dokument gut lesbar ist, bist du bereit für den Übungstest.")
    add_finish(doc, 'Gib dieses Arbeitsblatt zusammen mit der Bilddatei in deinem Ordner "IB" ab.')
    new_detached_workspace_section(doc, "A13", top_margin_cm=2.0, bottom_margin_cm=1.8)
    raw = [
        ("PROJEKTWOCHE BERN", 5),("12.–15. Mai 2027", 4),("ÜBER DIE WOCHE", 4),
        ("Vier Tage lang entdecken wir Bern in kleinen Gruppen. Wir besuchen bekannte Orte, lösen Aufgaben in der Altstadt und arbeiten an einem gemeinsamen Abschlussprodukt.", 6),
        ("Treffpunkt: Montag, 08.00 Uhr beim Schulhaus. Rückkehr: Donnerstag, ungefähr 17.30 Uhr.", 7),("BILDDATEI: a13_bern_altstadt.png", 8),
        ("MITBRINGEN", 4),("Trinkflasche", 2),("Regenjacke", 2),("Schreibzeug", 2),("kleiner Rucksack", 2),("Lunch für Montag", 7),
        ("TAGESABLAUF", 4),("Treffpunkt beim Schulhaus", 2),("Zugfahrt nach Bern", 2),("Zimmer beziehen und Material holen", 2),("Start mit der ersten Gruppenaufgabe", 8),
        ("PROGRAMM", 4),("Tag | Vormittag | Nachmittag | Abend", 2),("Montag | Anreise | Altstadt-Rallye | gemeinsames Kochen", 2),
        ("Dienstag | Bundeshaus | Gruppenauftrag | Spieleabend", 2),("Mittwoch | Museum | Freizeit | Abschluss vorbereiten", 2),("Donnerstag | Präsentationen | Rückreise | -", 8),
        ("WICHTIG", 4),("Wir bewegen uns in Bern grundsätzlich in Dreier- oder Vierergruppen. Treffpunkte und Zeiten müssen eingehalten werden. Bei Fragen meldet sich jede Gruppe bei der Lehrperson.", 6),
    ]
    for text, after in raw:
        p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(after); p.paragraph_format.line_spacing = 1
        x = p.add_run(text); set_run(x, size=11)
    finalise(doc, out, "A13 - Gesamtauftrag / Prüfungsvorbereitung")


def build(root: Path):
    sheets=root/'arbeitsblaetter'; assets=sheets/'assets'; assets.mkdir(parents=True,exist_ok=True)
    image=assets/'a13_bern_altstadt.png'; asset(image)
    build_document(sheets/'A13_Gesamtauftrag_Pruefungsvorbereitung.docx')


if __name__ == '__main__':
    from worksheet_runtime import run_single
    run_single('A13', build)
