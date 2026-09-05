from __future__ import annotations

import sys
from pathlib import Path

if __package__ in (None, ""):
    sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from PIL import Image, ImageDraw
from docx.shared import Pt

from course_common import (
    NAVY, TEAL_DARK, PALE_TEAL, WARM,
    base_doc, block, add_text, add_check, add_finish, finalise, set_run,
)
from course_build_helpers import clear_paragraph as _clear, new_detached_workspace_section


def asset(path: Path):
    # The actual scene is deliberately surrounded by generous empty margins.
    # This makes cropping a meaningful design decision in A12 instead of a
    # checkbox action on an already perfectly framed image.
    SW,H=1600,900; scene=Image.new('RGB',(SW,H),'#E7F1F4'); d=ImageDraw.Draw(scene)
    d.rectangle((0,0,SW,330),fill='#D9E9EF'); d.rectangle((0,330,SW,610),fill='#F1E6D4'); d.rectangle((0,610,SW,H),fill='#BED5C7')
    d.rectangle((360,310,1240,690),fill='#F4F0E7',outline='#17324D',width=7); d.polygon([(310,320),(800,95),(1290,320)],fill='#667D88',outline='#17324D')
    d.rectangle((720,500,880,690),fill='#B8D4D7',outline='#17324D',width=5)
    for y in (380,500):
        for x in (440,570,960,1090): d.rectangle((x,y,x+90,y+75),fill='#D8EEF1',outline='#17324D',width=4)
    for x in (180,1370): d.rectangle((x-15,480,x+15,700),fill='#7A5B43'); d.ellipse((x-115,330,x+115,560),fill='#4D7568')
    d.line((130,220,1470,220),fill='#445B65',width=4)
    for x in range(180,1450,120): d.line((x,220,x,252),fill='#445B65',width=3); d.ellipse((x-10,245,x+10,265),fill='#F5C86B',outline='#A67B31')
    for x in (270,1050): d.rectangle((x,710,x+250,742),fill='#8A6549'); d.line((x+40,742,x+15,800),fill='#8A6549',width=12); d.line((x+210,742,x+235,800),fill='#8A6549',width=12); d.rectangle((x-20,770,x+270,792),fill='#8A6549')
    d.rectangle((4,4,SW-5,H-5),outline='#D3DEE2',width=8)

    W,CH=1960,1200
    im=Image.new('RGB',(W,CH),'#F4F6F7')
    im.paste(scene,(180,160))
    im.save(path,quality=95)


def build_document(out: Path):
    doc=base_doc('A12','Selbstständig gestalten','Du entscheidest','Es gibt keine Zielvorlage mehr. Die Vorgaben sind klar – das Layout planst du selbst.','ein einseitiges Dokument mit bekannten Word-Werkzeugen selbstständig planen und übersichtlich gestalten.')
    t=block(doc,'01','FLYER'); r=t.cell(0,1); p=r.paragraphs[0]; _clear(p); x=p.add_run('Gestalte auf Seite 2 einen Flyer für den Schulhaus-Sommerabend'); set_run(x,size=12.5,bold=True,color=NAVY); add_text(r,'Alle Informationen stehen bereit. Du darfst die Reihenfolge der Bereiche verändern, aber keinen Inhalt weglassen.',9.35)
    t=block(doc,'PFLICHT',None,fill_left=WARM,fill_right=WARM,label_color=NAVY,label_size=9.3); r=t.cell(0,1); p=r.paragraphs[0]; _clear(p); x=p.add_run('Dein Flyer muss enthalten:'); set_run(x,size=9.7,bold=True,color=NAVY)
    for s in ['A4 Hochformat · alles auf einer Seite','Formatvorlagen für Titel und Überschriften','eine echte Aufzählung bei «MITBRINGEN»','eine echte Tabelle für das Programm','das Bild a12_sommerabend.png · zuschneiden · sinnvoll platzieren','Fusszeile «Sommerabend 2026» + automatische Seitenzahl']:
        add_text(r,'•  '+s,9.3,after=.25)
    t=block(doc,'DU','ENTSCHEIDEST',fill_left=PALE_TEAL,fill_right=PALE_TEAL,label_color=TEAL_DARK,label_size=13.5); p=t.cell(0,1).paragraphs[0]; _clear(p); x=p.add_run('Du entscheidest selbst über Anordnung, Grössen, Abstände, Bildposition und eine passende Farbe. '); set_run(x,size=9.35); x=p.add_run('Der Flyer soll vor allem schnell lesbar sein.'); set_run(x,size=9.35,bold=True,color=TEAL_DARK)
    add_check(doc,'Sind Datum, Zeit und Ort sofort sichtbar? Ist alles auf einer Seite? Sind Liste, Tabelle, Bild und Fusszeile wirklich vorhanden?')
    add_finish(doc,'Gib dieses Arbeitsblatt zusammen mit der Bilddatei in deinem Ordner "IB" ab.')
    new_detached_workspace_section(doc,'A12',top_margin_cm=1.8,bottom_margin_cm=1.8)
    raw=['SOMMERABEND IM SCHULHAUS','Freitag, 19. Juni 2026','18.00–21.00 Uhr','Schulhaus Sonnenberg · Innenhof','Wir beenden das Schuljahr gemeinsam mit Musik, Spielen, Essen und Zeit zum Zusammensitzen.','PROGRAMM','18.00 | Start und Musik','18.45 | Grill und Getränke','19.30 | Spielturnier','20.30 | Dessert und Abschluss','MITBRINGEN','Trinkbecher','Jacke oder Pullover','gute Laune','wer möchte: ein Kartenspiel','WICHTIG','Bei Regen findet der Sommerabend in der Aula statt. Die Teilnahme ist kostenlos.','BILDDATEI: a12_sommerabend.png']
    for text in raw:
        p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(4); x=p.add_run(text); set_run(x,size=11)
    finalise(doc,out,'A12 - Selbstständig gestalten')


def build(root: Path):
    sheets=root/'arbeitsblaetter'; assets=sheets/'assets'; assets.mkdir(parents=True,exist_ok=True)
    image=assets/'a12_sommerabend.png'; asset(image)
    build_document(sheets/'A12_Selbststaendig_gestalten.docx')


if __name__ == '__main__':
    from worksheet_runtime import run_single
    run_single('A12', build)
