from __future__ import annotations

import sys
from pathlib import Path

if __package__ in (None, ""):
    sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from PIL import Image, ImageDraw, ImageFont
from docx.shared import Pt

from build_runtime import resolve_font_paths
from course_common import (
    NAVY, TEAL_DARK,
    base_doc, block, add_text, add_step, add_tip, add_check, add_finish,
    add_picture, new_workspace_section, finalise, set_run,
)
from course_build_helpers import clear_paragraph as _clear


def preview(path: Path):
    reg,bold=resolve_font_paths(); W,H=1500,500
    im=Image.new('RGB',(W,H),'white'); d=ImageDraw.Draw(im)
    ft=ImageFont.truetype(bold,32); fh=ImageFont.truetype(bold,23); fb=ImageFont.truetype(reg,22)
    d.rounded_rectangle((20,20,W-20,H-20),radius=18,outline='#D3DEE2',width=3,fill='white')
    x0,y0=70,55; tw=1360; cols=[340,340,340,340]; line='#B8C4CA'
    d.rectangle((x0,y0,x0+tw,y0+60),fill='white',outline=line,width=2)
    s='SPORTTAG – ABLAUF'; box=d.textbbox((0,0),s,font=ft); d.text((x0+(tw-(box[2]-box[0]))/2,y0+11),s,font=ft,fill='#17324D')
    rows=[['Zeit','Ort','Aktivität','Gruppe'],['09.00','Aula','Begrüssung','Alle'],['09.30','Turnhalle','Staffellauf','A'],['09.30','Sportplatz','Fussball','B'],['10.30','Aula','Pause','Alle']]
    y=y0+60
    for ri,row in enumerate(rows):
        x=x0
        for j,(w,val) in enumerate(zip(cols,row)):
            d.rectangle((x,y,x+w,y+54),fill='white',outline=line,width=2)
            f=fh if ri==0 else fb; c='#17324D'
            box=d.textbbox((0,0),val,font=f)
            tx=x+(w-(box[2]-box[0]))/2 if ri==0 or j in (0,3) else x+14
            d.text((tx,y+13),val,font=f,fill=c); x+=w
        y+=54
    im.save(path,quality=95)


def table_tools_anchor(path: Path):
    reg,bold=resolve_font_paths(); W,H=1500,280
    im=Image.new('RGB',(W,H),'white'); d=ImageDraw.Draw(im)
    fl=ImageFont.truetype(bold,23); fs=ImageFont.truetype(reg,20); fb=ImageFont.truetype(bold,21)
    d.rounded_rectangle((18,18,W-18,H-18),radius=18,outline='#D3DEE2',width=3,fill='white')
    panels=[
        (55,'Einfügen → Tabelle','4 × 5'),
        (535,'Tabellenlayout','Darüber einfügen'),
        (1015,'Tabellenlayout','Zellen verbinden'),
    ]
    for x,title,action in panels:
        d.text((x,38),title,font=fl,fill='#1D6765')
        d.rounded_rectangle((x,86,x+410,205),radius=10,outline='#AEBCC3',width=2,fill='#F8FAFB')
        if action=='4 × 5':
            gx,gy=x+28,104
            for r in range(5):
                for c in range(4):
                    d.rectangle((gx+c*45,gy+r*17,gx+c*45+40,gy+r*17+13),fill='#EAF4F3',outline='#237B78',width=1)
            d.text((x+245,124),'4 × 5',font=fb,fill='#17324D')
        elif action=='Darüber einfügen':
            d.line((x+48,145,x+180,145),fill='#17324D',width=4)
            d.polygon([(x+108,100),(x+88,128),(x+128,128)],fill='#237B78')
            d.text((x+210,122),action,font=fs,fill='#17324D')
        else:
            for c in range(4):
                d.rectangle((x+35+c*44,112,x+75+c*44,154),outline='#17324D',width=2)
            d.line((x+75,112,x+75,154),fill='#F8FAFB',width=5)
            d.line((x+119,112,x+119,154),fill='#F8FAFB',width=5)
            d.line((x+163,112,x+163,154),fill='#F8FAFB',width=5)
            d.text((x+235,122),action,font=fs,fill='#17324D')
        d.text((x,222),action,font=fs,fill='#5E6D78')
    im.save(path,quality=95)


def build_document(out: Path, preview_path: Path, tools_path: Path):
    doc=base_doc('A8','Tabellen','Infos ins Raster bringen','Du baust eine einfache Tabelle und passt sie Schritt für Schritt an.','eine Tabelle erstellen, Daten in Zellen eintragen, eine Zeile ergänzen und Zellen verbinden.')
    t=block(doc,'01','AUFGABE'); r=t.cell(0,1); p=r.paragraphs[0]; _clear(p); x=p.add_run('Baue den Sporttag-Plan nach'); set_run(x,size=12.8,bold=True,color=NAVY)
    add_text(r,'Arbeite auf Seite 2. Die Daten stehen dort bereits bereit.',9.5,after=.2)
    add_text(r,'NEU: Einfügen → Tabelle → 4 × 5.  Neue Zeile → Tabellenlayout → Darüber einfügen.  Verbinden → Tabellenlayout → Zellen verbinden.',9.1,bold=True,color=TEAL_DARK,after=.1)
    add_picture(r,tools_path,10.6)
    add_picture(r,preview_path,10.9)
    for letter,parts in [
        ('A',[('Unter «HIER TABELLE EINFÜGEN»',True,False,NAVY),('  →  Tabelle mit 4 Spalten und 5 Zeilen einfügen',False,False,None)]),
        ('B',[('Kopfzeile + vier Datenzeilen',True,False,NAVY),('  →  Angaben von oben in die richtigen Zellen übertragen',False,False,None)]),
        ('C',[('Kopfzeile',True,False,NAVY),('  →  fett und zentriert',False,False,None)]),
        ('D',[('Zeit und Gruppe',True,False,NAVY),('  →  zentrieren',False,False,None)]),
        ('E',[('Neue Zeile über der Tabelle einfügen',True,False,NAVY),('  →  alle 4 Zellen dieser Zeile verbinden',False,False,None)]),
        ('F',[('In die verbundene Zelle',True,False,NAVY),('  →  «SPORTTAG – ABLAUF», 16 pt, fett, zentriert',False,False,None)]),
    ]: add_step(r,letter,parts)
    add_tip(doc,'Die Vorlage zeigt nur die verlangten Kernfunktionen. Spaltenbreiten, Farben und besondere Tabellen-Designs musst du nicht nachbauen.')
    add_check(doc,'Hat deine Tabelle 4 Spalten? Sind alle Daten in der richtigen Zeile? Geht der Titel über die ganze Tabelle?')
    add_finish(doc)
    new_workspace_section(doc,'A8')
    p=doc.add_paragraph(); x=p.add_run('DATEN FÜR DIE TABELLE'); set_run(x,size=16,bold=True,color=NAVY)
    for s in ['Zeit | Ort | Aktivität | Gruppe','09.00 | Aula | Begrüssung | Alle','09.30 | Turnhalle | Staffellauf | A','09.30 | Sportplatz | Fussball | B','10.30 | Aula | Pause | Alle']:
        p=doc.add_paragraph(); x=p.add_run(s); set_run(x,size=11)
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(12); x=p.add_run('HIER TABELLE EINFÜGEN'); set_run(x,size=12,bold=True,color=TEAL_DARK)
    finalise(doc,out,'A8 - Tabellen')


def build(root: Path):
    sheets=root/'arbeitsblaetter'; prev=sheets/'assets'/'vorlagen'; prev.mkdir(parents=True,exist_ok=True)
    image=prev/'a8_sporttag_tabelle_vorlage.png'; tools=prev/'a8_tabellen_werkzeuge.png'
    preview(image); table_tools_anchor(tools)
    build_document(sheets/'A8_Tabellen.docx',image,tools)


if __name__ == '__main__':
    from worksheet_runtime import run_single
    run_single('A8', build)
