from __future__ import annotations

import sys
from pathlib import Path

if __package__ in (None, ""):
    sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from PIL import Image, ImageDraw, ImageFont
from docx.shared import Pt
from docx.enum.text import WD_BREAK

from build_runtime import resolve_font_paths
from course_common import (
    NAVY, TEAL,
    base_doc, block, add_text, add_step, add_tip, add_check, add_finish,
    add_picture, finalise, set_run,
)
from course_build_helpers import clear_paragraph as _clear, new_detached_workspace_section


def preview(path: Path):
    reg, bold = resolve_font_paths()
    W, H = 1500, 440
    im = Image.new('RGB', (W, H), 'white')
    d = ImageDraw.Draw(im)
    f_small = ImageFont.truetype(reg, 24)
    f_small_b = ImageFont.truetype(bold, 24)
    f_label = ImageFont.truetype(bold, 18)
    f_num = ImageFont.truetype(bold, 25)
    line = '#D3DEE2'
    pale = '#F3F6F7'
    pale_teal = '#EAF4F3'
    d.rounded_rectangle((18, 18, W-18, H-18), radius=18, outline=line, width=3, fill='white')
    for idx, x in enumerate((120, 800), start=1):
        y = 62
        pw, ph = 520, 320
        d.text((x, 24), f'Übungsseite {idx}', font=f_small_b, fill='#237B78')
        d.rectangle((x, y, x+pw, y+ph), fill='white', outline='#9FB0B7', width=3)
        d.rectangle((x+3, y+3, x+pw-3, y+55), fill=pale_teal)
        d.rectangle((x+3, y+ph-52, x+pw-3, y+ph-3), fill=pale)
        d.line((x+30, y+55, x+pw-30, y+55), fill=line, width=2)
        d.line((x+30, y+ph-52, x+pw-30, y+ph-52), fill=line, width=2)
        d.text((x+34, y+17), 'SCHULAUSFLUG LUZERN', font=f_small_b, fill='#17324D')
        d.text((x+40, y+92), 'Reisebericht', font=f_small_b, fill='#237B78')
        d.text((x+40, y+135), 'Der Text steht im normalen Seitenbereich.', font=f_small, fill='#5E6D78')
        footer = f'Seite {idx}'
        bbox = d.textbbox((0, 0), footer, font=f_num)
        d.text((x+(pw-(bbox[2]-bbox[0]))/2, y+ph-42), footer, font=f_num, fill='#17324D')
        if idx == 1:
            d.text((22, y+18), 'Kopfzeile', font=f_label, fill='#1D6765')
            d.line((95, y+31, x-6, y+31), fill='#237B78', width=2)
            d.text((22, y+ph-38), 'Fusszeile', font=f_label, fill='#667684')
            d.line((92, y+ph-26, x-6, y+ph-26), fill='#9FB0B7', width=2)
    im.save(path, quality=95)


def build_document(out: Path, preview_path: Path):
    doc = base_doc(
        'A10', 'Kopf- & Fusszeile', 'Infos auf jeder Seite',
        'Du setzt Informationen in den oberen und unteren Seitenbereich.',
        'eine Kopfzeile, eine Fusszeile und automatische Seitenzahlen einfügen.'
    )
    t = block(doc, '01', 'ÜBUNGSSEITEN')
    r = t.cell(0,1)
    p = r.paragraphs[0]
    _clear(p)
    x = p.add_run('Ergänze den Reisebericht')
    set_run(x, size=12.8, bold=True, color=NAVY)
    add_text(r, 'Arbeite nur im Reisebericht auf den beiden folgenden Übungsseiten. Der normale Text ist bereits fertig.', 9.5, after=.6)
    add_picture(r, preview_path, 10.9)
    add_step(r, 'A', [('Doppelklicke ganz oben auf der ersten Übungsseite', True, False, NAVY), ('  →  Kopfzeile öffnen', False, False, None)])
    add_step(r, 'B', [('Kopfzeile', True, False, NAVY), ('  →  «SCHULAUSFLUG LUZERN» eingeben', False, False, None)])
    add_step(r, 'C', [('Doppelklicke ganz unten auf der ersten Übungsseite', True, False, NAVY), ('  →  Fusszeile öffnen', False, False, None)])
    add_step(r, 'D', [('Fusszeile', True, False, NAVY), ('  →  «Seite » eingeben', False, False, None)])
    add_step(r, 'E', [('Direkt hinter «Seite »', True, False, NAVY), ('  →  Einfügen → Seitenzahl → Aktuelle Position → eine einfache Zahl ohne Rahmen/Design wählen', False, False, None)])
    add_step(r, 'F', [('Kontrolliere die zweite Übungsseite', True, False, NAVY), ('  →  Kopf- und Fusszeile erscheinen dort automatisch; die Zahl steigt von 1 auf 2.', False, False, None)])
    add_tip(doc, 'Kopfzeile = Bereich oben. Fusszeile = Bereich unten. Beide wiederholen sich auf den Übungsseiten.')
    add_tip(doc, 'Die Seitenzahl ist automatisch: Tippe nur «Seite » und wähle Einfügen → Seitenzahl → Aktuelle Position → eine einfache Zahl ohne Rahmen/Design. Die Zahl nicht von Hand schreiben.', 'MERKE')
    add_check(doc, 'Oben steht auf beiden Übungsseiten «SCHULAUSFLUG LUZERN». Unten steht auf der ersten Übungsseite «Seite 1» und auf der zweiten «Seite 2».')
    add_finish(doc)
    new_detached_workspace_section(
        doc,
        top_margin_cm=2.2,
        bottom_margin_cm=2.0,
        left_margin_cm=2.0,
        right_margin_cm=2.0,
        header_distance_cm=.7,
        footer_distance_cm=.7,
    )
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4); x = p.add_run('REISEBERICHT LUZERN'); set_run(x, size=20, bold=True, color=NAVY)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(5); x = p.add_run('Der Morgen'); set_run(x, size=15, bold=True, color=TEAL)
    for s in [
        'Um 08.00 Uhr treffen wir uns beim Schulhaus. Gemeinsam fahren wir mit dem Zug nach Luzern.',
        'Nach der Ankunft gehen wir direkt zum Verkehrshaus. Dort arbeiten wir in kleinen Gruppen.',
        'Jede Gruppe sucht drei Ausstellungsstücke und notiert dazu die wichtigsten Informationen.',
    ]:
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(8); x = p.add_run(s); set_run(x, size=11)
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(5); x = p.add_run('Im Verkehrshaus'); set_run(x, size=15,bold=True,color=TEAL)
    for s in [
        'Besonders spannend sind die Bereiche Luftfahrt und Raumfahrt. Einige Gruppen besuchen zusätzlich die Eisenbahn-Ausstellung.',
        'Vor dem Mittag treffen wir uns wieder beim Haupteingang und vergleichen unsere Notizen.',
    ]:
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(8); x = p.add_run(s); set_run(x, size=11)
    p = doc.add_paragraph(); p.add_run().add_break(WD_BREAK.PAGE)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(5); x = p.add_run('Mittag und Altstadt'); set_run(x, size=15,bold=True,color=TEAL)
    for s in [
        'Das Mittagessen verbringen wir am See. Danach laufen wir gemeinsam in Richtung Altstadt.',
        'Beim Rundgang sehen wir die Kapellbrücke und verschiedene historische Gebäude. Anschliessend bleibt Zeit für eine kurze Pause.',
    ]:
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(8); x = p.add_run(s); set_run(x, size=11)
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(5); x = p.add_run('Rückfahrt'); set_run(x, size=15,bold=True,color=TEAL)
    for s in [
        'Um 16.30 Uhr nehmen wir den Zug zurück. Die Ankunft beim Schulhaus ist ungefähr um 18.00 Uhr.',
        'Damit endet unser Schulausflug nach Luzern.',
    ]:
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(8); x = p.add_run(s); set_run(x, size=11)
    finalise(doc, out, 'A10 - Kopf- und Fusszeile / Seitenzahlen')


def build(root: Path):
    sheets=root/'arbeitsblaetter'; prev=sheets/'assets'/'vorlagen'; prev.mkdir(parents=True,exist_ok=True)
    image=prev/'a10_kopf_fuss_vorlage.png'; preview(image)
    build_document(sheets/'A10_Kopf_Fusszeile_Seitenzahlen.docx',image)


if __name__ == '__main__':
    from worksheet_runtime import run_single
    run_single('A10', build)
