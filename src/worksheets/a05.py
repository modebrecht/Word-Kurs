from __future__ import annotations

import sys
from pathlib import Path

if __package__ in (None, ""):
    sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from PIL import Image, ImageDraw, ImageFont
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from build_runtime import resolve_font_paths
from course_common import (
    NAVY, TEAL, TEAL_DARK, PALE,
    base_doc, block, add_text, add_step, add_tip, add_picture, finalise, set_run,
)
from course_build_helpers import clear_paragraph as _clear


def school_icon(path: Path):
    W,H=420,300
    im=Image.new('RGBA',(W,H),(255,255,255,0)); d=ImageDraw.Draw(im)
    d.rounded_rectangle((55,95,365,260),radius=12,fill='#EAF4F3',outline='#237B78',width=5)
    d.polygon([(45,105),(210,35),(375,105)],fill='#237B78')
    d.rounded_rectangle((180,170,240,260),radius=6,fill='#17324D')
    for x in (95,275): d.rounded_rectangle((x,145,x+55,195),radius=5,fill='white',outline='#17324D',width=4)
    d.line((210,35,210,2),fill='#17324D',width=5); d.polygon([(210,4),(274,20),(210,36)],fill='#17324D')
    im.save(path)


def preview(path: Path, icon_path: Path):
    reg,bold=resolve_font_paths(); W,H=1500,520
    im=Image.new('RGB',(W,H),'white'); d=ImageDraw.Draw(im)
    ft=ImageFont.truetype(bold,48); fs=ImageFont.truetype(bold,26); fh=ImageFont.truetype(bold,24); fb=ImageFont.truetype(reg,23)
    d.rounded_rectangle((18,18,W-18,H-18),radius=18,outline='#D3DEE2',width=3,fill='white')

    icon=Image.open(icon_path).convert('RGBA'); icon.thumbnail((140,100)); im.paste(icon,((W-icon.width)//2,14),icon)
    s='SCHULFEST 2026'; box=d.textbbox((0,0),s,font=ft); d.text(((W-(box[2]-box[0]))/2,118),s,font=ft,fill='#17324D')
    s='Freitag, 12. Juni · 17.30 Uhr'; box=d.textbbox((0,0),s,font=fs); d.text(((W-(box[2]-box[0]))/2,176),s,font=fs,fill='#172A3A')

    x=115; y=235
    d.text((x,y),'Mitnehmen',font=fh,fill='#172A3A'); y+=36
    for s in ['Trinkflasche','Jacke für den Abend','gute Laune']:
        d.text((x+8,y),'•',font=fb,fill='#172A3A'); d.text((x+44,y),s,font=fb,fill='#172A3A'); y+=32

    y=365
    d.text((x,y),'Ablauf',font=fh,fill='#172A3A'); y+=34
    for i,s in enumerate(['Begrüssung','Spiel & Essen','Musik in der Aula'],1):
        d.text((x+5,y),f'{i}.',font=fb,fill='#172A3A'); d.text((x+48,y),s,font=fb,fill='#172A3A'); y+=30
    im.save(path,quality=95)


def build_document(out: Path, preview_path: Path, icon: Path):
    doc=base_doc('A5','Dokument retten','Rette das Chaos-Dokument','Vergleiche mit der Vorlage und repariere gezielte Formatierungsfehler.','bekannte Formatierungsfehler erkennen und reparieren. Eine vorhandene Grafik kann ich mit sichtbarer Hilfe verkleinern und an die richtige Stelle verschieben.')
    t=block(doc,'01','REPARIEREN'); r=t.cell(0,1); p=r.paragraphs[0]; _clear(p); x=p.add_run('Mach das Chaos wieder ordentlich'); set_run(x,size=12.8,bold=True,color=NAVY); add_text(r,'Vergleiche immer wieder mit der Vorlage. Arbeite von A bis F.',9.2); add_picture(r,preview_path,9.7)
    add_step(r,'A',[('Titel «SCHULFEST 2026»',True,False,NAVY),('  →  Arial, 20 pt, fett, dunkelblau, zentriert',False,False,None)])
    add_step(r,'B',[('Datum',True,False,NAVY),('  →  Arial, 11 pt, fett, zentriert',False,False,None)])
    add_step(r,'C',[('«Mitnehmen» + drei Dinge',True,False,NAVY),('  →  Überschrift fett, Dinge als echte Word-Aufzählung',False,False,None)])
    add_step(r,'D',[('«Ablauf» + drei Schritte',True,False,NAVY),('  →  Überschrift fett, Schritte als echte Word-Nummerierung',False,False,None)])
    add_step(r,'E',[('Text in den beiden Listen',True,False,NAVY),('  →  Arial, 11 pt, links, Zeilenabstand 1,0',False,False,None)])
    add_step(r,'F',[('MINI-NEU · Grafik',True,False,TEAL_DARK),('  →  anklicken → Bildformat → Grösse → Breite ca. 2,2 cm; danach über den Titel verschieben und zentrieren',False,False,None)])
    t=block(doc,'CHAOS','DOKUMENT 01',fill_left=PALE,fill_right='FBFCFC',label_color=TEAL_DARK,label_size=11.0); r=t.cell(0,1); p=r.paragraphs[0]; _clear(p)
    x=p.add_run('SCHULFEST 2026'); set_run(x,name='Times New Roman',size=24,italic=True,color='C05A2B')
    p=r.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.RIGHT; x=p.add_run('Freitag, 12. Juni · 17.30 Uhr'); set_run(x,name='Courier New',size=9,color='7A2B83')
    p=r.add_paragraph(); p.paragraph_format.space_after=Pt(0); p.add_run().add_picture(str(icon),width=Cm(3.3))
    for text,name,size,bold,italic,color,align in [
        ('Mitnehmen','Arial',13,False,True,'C05A2B',WD_ALIGN_PARAGRAPH.RIGHT),
        ('1. Trinkflasche','Comic Sans MS',11,False,False,'333333',WD_ALIGN_PARAGRAPH.LEFT),
        ('2. Jacke für den Abend','Comic Sans MS',11,False,False,'333333',WD_ALIGN_PARAGRAPH.LEFT),
        ('3. gute Laune','Comic Sans MS',11,False,False,'333333',WD_ALIGN_PARAGRAPH.LEFT),
        ('Ablauf','Arial',13,False,False,TEAL,WD_ALIGN_PARAGRAPH.CENTER),
        ('• Begrüssung','Times New Roman',11,False,False,'333333',WD_ALIGN_PARAGRAPH.LEFT),
        ('• Spiel & Essen','Times New Roman',11,False,False,'333333',WD_ALIGN_PARAGRAPH.LEFT),
        ('• Musik in der Aula','Times New Roman',11,False,False,'333333',WD_ALIGN_PARAGRAPH.LEFT),
    ]:
        p=r.add_paragraph(); p.alignment=align; p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=.9; x=p.add_run(text); set_run(x,name=name,size=(10 if size>=11 else size),bold=bold,italic=italic,color=color)
    add_tip(doc,'Grafik verschieben: anklicken → Ctrl + X → Cursor über den Titel → Ctrl + V. Danach zentrieren. Die Breite stellst du wie in Schritt F unter Bildformat → Grösse ein.  FERTIG? Gib das Blatt in deinem Ordner "IB" ab.')
    finalise(doc,out,'A5 - Rette das Chaos-Dokument')


def build(root: Path):
    sheets=root/'arbeitsblaetter'; assets=sheets/'assets'; prev=assets/'vorlagen'; assets.mkdir(parents=True,exist_ok=True); prev.mkdir(parents=True,exist_ok=True)
    icon=assets/'a5_school_icon.png'; image=prev/'a5_schulfest_vorlage.png'
    school_icon(icon); preview(image,icon)
    build_document(sheets/'A5_Rette_das_Chaos_Dokument.docx',image,icon)


if __name__ == '__main__':
    from worksheet_runtime import run_single
    run_single('A5', build)
