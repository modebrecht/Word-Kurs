from __future__ import annotations

import sys
from pathlib import Path

if __package__ in (None, ""):
    sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from PIL import Image, ImageDraw, ImageFont
from docx.shared import Pt

from build_runtime import resolve_font_paths
from course_common import (
    NAVY,
    base_doc, block, add_text, add_step, add_tip, add_check, add_finish,
    add_picture, new_workspace_section, finalise, set_run,
)
from course_build_helpers import clear_paragraph as _clear


def preview(path: Path):
    reg,bold=resolve_font_paths(); W,H=1500,520
    im=Image.new('RGB',(W,H),'white'); d=ImageDraw.Draw(im)
    ft=ImageFont.truetype(bold,46); f1=ImageFont.truetype(bold,30); f2=ImageFont.truetype(bold,24); fb=ImageFont.truetype(reg,22)
    d.rounded_rectangle((18,18,W-18,H-18),radius=18,outline='#D3DEE2',width=3,fill='white')
    d.text((65,40),'UNSER SCHULAUSFLUG NACH LUZERN',font=ft,fill='#17324D')
    d.text((65,125),'Der Morgen',font=f1,fill='#237B78'); d.text((65,170),'Treffpunkt',font=f2,fill='#17324D'); d.text((65,205),'Wir treffen uns um 08.00 Uhr beim Schulhaus und fahren gemeinsam zum Bahnhof.',font=fb,fill='#5E6D78')
    d.text((65,270),'Im Verkehrshaus',font=f1,fill='#237B78'); d.text((65,315),'Unsere Aufgabe',font=f2,fill='#17324D'); d.text((65,350),'In Gruppen suchen wir drei Ausstellungsstücke und notieren die wichtigsten Informationen.',font=fb,fill='#5E6D78')
    d.text((65,415),'Mittag und Altstadt',font=f1,fill='#237B78'); d.text((65,458),'Nach dem Mittagessen spazieren wir gemeinsam durch die Altstadt.',font=fb,fill='#5E6D78')
    im.save(path,quality=95)


def build_document(out: Path, preview_path: Path):
    doc=base_doc('A9','Formatvorlagen','Überschriften mit System','Du gibst Textteilen eine feste Rolle, statt jede Überschrift von Hand zu gestalten.','Titel, Überschrift 1, Überschrift 2 und normalen Text mit den passenden Word-Formatvorlagen auszeichnen.')
    t=block(doc,'01','SEITE 2'); r=t.cell(0,1); p=r.paragraphs[0]; _clear(p); x=p.add_run('Gib jedem Textteil die richtige Rolle'); set_run(x,size=12.8,bold=True,color=NAVY); add_text(r,'Arbeite auf Seite 2. Ändere Schriftgrösse oder Fett nicht von Hand.',9.5); add_picture(r,preview_path,10.9)
    add_step(r,'A',[('«UNSER SCHULAUSFLUG NACH LUZERN»',True,False,NAVY),('  →  Formatvorlage «Titel»',False,False,None)])
    add_step(r,'B',[('«Der Morgen», «Im Verkehrshaus», «Mittag und Altstadt»',True,False,NAVY),('  →  «Überschrift 1»',False,False,None)])
    add_step(r,'C',[('«Treffpunkt», «Unsere Aufgabe», «Rückfahrt»',True,False,NAVY),('  →  «Überschrift 2»',False,False,None)])
    add_step(r,'D',[('Alle übrigen Absätze',True,False,NAVY),('  →  «Standard»',False,False,None)])
    add_tip(doc,'Titel  >  Überschrift 1  >  Überschrift 2  >  Standard','MERKE')
    add_tip(doc,'Formatvorlagen findest du bei «Start». Klicke zuerst in den Absatz und dann auf die passende Formatvorlage.')
    add_check(doc,'Sehen alle Überschriften derselben Stufe gleich aus? Dann hast du die Formatvorlagen richtig eingesetzt.')
    add_finish(doc)
    new_workspace_section(doc,'A9')
    raw=[
        'UNSER SCHULAUSFLUG NACH LUZERN','Der Morgen','Treffpunkt','Wir treffen uns um 08.00 Uhr beim Schulhaus und fahren gemeinsam zum Bahnhof. Im Zug sitzen wir in unseren Gruppen.','Im Verkehrshaus','Unsere Aufgabe','In Gruppen suchen wir drei Ausstellungsstücke. Zu jedem Stück notieren wir zwei wichtige Informationen.','Mittag und Altstadt','Nach dem Mittagessen spazieren wir gemeinsam durch die Altstadt. Danach bleibt noch Zeit für eine kurze Pause am See.','Rückfahrt','Um 16.30 Uhr fahren wir zurück. Die Ankunft beim Schulhaus ist ungefähr um 18.00 Uhr.'
    ]
    for s in raw:
        p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(6); x=p.add_run(s); set_run(x,size=11)
    finalise(doc,out,'A9 - Formatvorlagen und Überschriften')


def build(root: Path):
    sheets=root/'arbeitsblaetter'; prev=sheets/'assets'/'vorlagen'; prev.mkdir(parents=True,exist_ok=True)
    image=prev/'a9_formatvorlagen_vorlage.png'; preview(image)
    build_document(sheets/'A9_Formatvorlagen_und_Ueberschriften.docx',image)


if __name__ == '__main__':
    from worksheet_runtime import run_single
    run_single('A9', build)
