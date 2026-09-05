from __future__ import annotations

import sys
from pathlib import Path

if __package__ in (None, ""):
    sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from PIL import Image, ImageDraw, ImageFont
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

from build_runtime import resolve_font_paths
from course_common import (
    NAVY, TEAL, TEAL_DARK,
    base_doc, block, add_text, add_step, add_tip, add_check, add_finish,
    add_picture, new_workspace_section, finalise, set_run,
)
from course_build_helpers import clear_paragraph as _clear


def layout_tools_preview(path: Path):
    reg, bold = resolve_font_paths()
    W, H = 1500, 320
    im = Image.new('RGB', (W, H), 'white')
    d = ImageDraw.Draw(im)
    f_tab = ImageFont.truetype(bold, 30)
    f_label = ImageFont.truetype(bold, 27)
    f_small = ImageFont.truetype(reg, 22)
    f_key = ImageFont.truetype(bold, 24)
    border = '#D3DEE2'
    pale = '#F3F6F7'
    accent = '#237B78'
    navy = '#17324D'

    d.rounded_rectangle((16, 16, W-16, H-16), radius=18, outline=border, width=3, fill='white')
    d.rectangle((16, 16, W-16, 80), fill=pale)
    for x, label in [(48, 'Start'), (170, 'Einfügen'), (325, 'Zeichnen'), (480, 'Entwurf')]:
        d.text((x, 36), label, font=f_small, fill='#667684')
    d.text((625, 32), 'Layout', font=f_tab, fill=navy)
    d.line((622, 75, 735, 75), fill=accent, width=5)

    # Orientation card
    d.rounded_rectangle((65, 112, 600, 262), radius=12, outline=accent, width=4, fill='white')
    d.text((92, 128), 'Ausrichtung', font=f_small, fill='#667684')
    # portrait/landscape page icons
    d.rectangle((100, 166, 148, 232), outline=navy, width=4)
    d.rectangle((180, 181, 250, 229), outline=accent, width=4)
    d.text((285, 175), 'Querformat', font=f_label, fill=navy)

    # Margins card
    d.rounded_rectangle((640, 112, 1175, 262), radius=12, outline=accent, width=4, fill='white')
    d.text((668, 128), 'Seitenränder', font=f_small, fill='#667684')
    d.rectangle((680, 164, 744, 236), outline=navy, width=4)
    d.rectangle((694, 176, 730, 224), outline=accent, width=3)
    d.text((785, 175), 'Schmal', font=f_label, fill=navy)

    # Page-break keyboard anchor; not a finished page preview.
    d.text((1210, 125), 'Neue Seite', font=f_small, fill='#667684')
    d.rounded_rectangle((1210, 164, 1435, 222), radius=10, outline='#9FB0B7', width=3, fill=pale)
    d.text((1240, 180), 'Ctrl + Enter', font=f_key, fill=navy)

    path.parent.mkdir(parents=True, exist_ok=True)
    im.save(path, quality=95)


def _add_true_bullet(doc, text: str):
    p=doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after=Pt(2)
    x=p.add_run(text); set_run(x,size=11)
    return p


def build_document(out: Path, preview_path: Path):
    doc=base_doc('A6','Seitenlayout','Die Seite passend einstellen','Du arbeitest heute nur am Aufbau der Seite – nicht am Text.','eine Seite auf Querformat stellen, Seitenränder ändern und mit einem Seitenumbruch gezielt eine neue Seite beginnen.')
    t=block(doc,'01','SEITENLAYOUT'); r=t.cell(0,1); p=r.paragraphs[0]; _clear(p); x=p.add_run('Bearbeite nur den Übungsteil'); set_run(x,size=12.8,bold=True,color=NAVY)
    add_text(r,'Klicke zuerst irgendwo auf Seite 2. Dort beginnt der Übungsteil; der Text ist bereits richtig formatiert.',9.5,after=.35)
    add_text(r,'NEU: Querformat → Layout → Ausrichtung → Querformat.  Seitenränder → Layout → Seitenränder → Schmal.',9.15,bold=True,color=TEAL_DARK,after=.35)
    add_picture(r, preview_path, 11.0)
    add_step(r,'A',[('Übungsteil ab Seite 2',True,False,NAVY),('  →  Querformat',False,False,None)])
    add_step(r,'B',[('Übungsteil ab Seite 2',True,False,NAVY),('  →  Seitenränder «Schmal»',False,False,None)])
    add_step(r,'C',[('Klicke direkt vor «PACKLISTE»',True,False,NAVY),('  →  Ctrl + Enter',False,False,None)])
    add_step(r,'D',[('Kontrolliere',True,False,NAVY),('  →  «PACKLISTE» beginnt jetzt auf einer neuen Seite.',False,False,None)])
    add_tip(doc,'Hochformat = ▯   |   Querformat = ▭   |   Neue Seite = Ctrl + Enter','MERKE')
    add_tip(doc,'Für eine neue Seite nicht viele Male Enter drücken. Ein Seitenumbruch bleibt auch dann richtig, wenn später Text dazukommt.')
    add_check(doc,'Seite 1 bleibt Hochformat. Die Übungsseiten sind Querformat und haben schmale Ränder. «PACKLISTE» beginnt auf einer neuen Seite.')
    add_finish(doc)
    sec=new_workspace_section(doc,'A6'); sec.orientation=WD_ORIENT.PORTRAIT; sec.page_width=Cm(21); sec.page_height=Cm(29.7)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; x=p.add_run('SCHULAUSFLUG NACH LUZERN'); set_run(x,size=20,bold=True,color=NAVY)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(12); x=p.add_run('Dienstag, 22. September'); set_run(x,size=11,bold=True,color=TEAL)
    p=doc.add_paragraph(); x=p.add_run('PROGRAMM'); set_run(x,size=13,bold=True,color=NAVY)
    for tm,txt in [('08.00 Uhr','Treffpunkt beim Schulhaus'),('08.20 Uhr','Abfahrt mit dem Zug'),('10.00 Uhr','Verkehrshaus'),('12.15 Uhr','Mittagspause am See'),('14.00 Uhr','Altstadt-Rundgang'),('16.30 Uhr','Rückfahrt')]:
        p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(4); x=p.add_run(tm+'  –  '); set_run(x,size=11,bold=True,color=NAVY); x=p.add_run(txt); set_run(x,size=11)
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(14); x=p.add_run('PACKLISTE'); set_run(x,size=13,bold=True,color=NAVY)
    for item in ['Trinkflasche','Lunch','Regenjacke','Schreibzeug','Halbtax / Billet falls vorhanden']:
        _add_true_bullet(doc,item)
    finalise(doc,out,'A6 - Seitenlayout')


def build(root: Path):
    sheets=root/'arbeitsblaetter'; prev=sheets/'assets'/'vorlagen'; prev.mkdir(parents=True, exist_ok=True)
    image=prev/'a6_layout_werkzeuge.png'; layout_tools_preview(image)
    build_document(sheets/'A6_Seitenlayout.docx', image)


if __name__ == '__main__':
    from worksheet_runtime import run_single
    run_single('A6', build)
