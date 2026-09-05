from __future__ import annotations

import argparse
import re
import zipfile
from pathlib import Path
from xml.etree import ElementTree

from docx import Document
from PIL import Image

from grading import effort_grade, final_grade_with_drop_str, swiss_grade_str

EXPECTED_DOCX = (
    "A1_Text_formatieren.docx",
    "A2_Nach_Vorlage_gestalten.docx",
    "A3_Absaetze_und_Ordnung.docx",
    "A4_Listen_und_Nummerierungen.docx",
    "A5_Rette_das_Chaos_Dokument.docx",
    "A6_Seitenlayout.docx",
    "A7_Bilder_in_Word.docx",
    "A8_Tabellen.docx",
    "A9_Formatvorlagen_und_Ueberschriften.docx",
    "A10_Kopf_Fusszeile_Seitenzahlen.docx",
    "A11_Dokument_nach_Vorlage_nachbauen.docx",
    "A12_Selbststaendig_gestalten.docx",
    "A13_Gesamtauftrag_Pruefungsvorbereitung.docx",
    "Uebungstest_Word.docx",
    "Uebungstest_Ausgangsdokument.docx",
    "Benoteter_Steckbrief.docx",
    "Word_Test.docx",
    "Word_Test_Ausgangsdokument.docx",
    "Word_Test_Korrekturblatt.docx",
)

EXPECTED_ASSETS = (
    "assets/a7_schulhaus.png",
    "assets/a11_klassenlager_berge.png",
    "assets/a12_sommerabend.png",
    "assets/a13_bern_altstadt.png",
    "assets/uebungstest_greifensee.png",
    "assets/word_test_rheinfall.png",
)

EXPECTED_PACKAGES = {
    "pakete/A7_Bilder_in_Word.zip": (
        "A7_Bilder_in_Word.docx",
        "a7_schulhaus.png",
    ),
    "pakete/A11_Dokument_nach_Vorlage_nachbauen.zip": (
        "A11_Dokument_nach_Vorlage_nachbauen.docx",
        "a11_klassenlager_berge.png",
    ),
    "pakete/A12_Selbststaendig_gestalten.zip": (
        "A12_Selbststaendig_gestalten.docx",
        "a12_sommerabend.png",
    ),
    "pakete/A13_Gesamtauftrag_Pruefungsvorbereitung.zip": (
        "A13_Gesamtauftrag_Pruefungsvorbereitung.docx",
        "a13_bern_altstadt.png",
    ),
    "pakete/Uebungstest_Word_Paket.zip": (
        "Uebungstest_Word.docx",
        "Uebungstest_Ausgangsdokument.docx",
        "uebungstest_greifensee.png",
    ),
    "pakete/Word_Test_Paket.zip": (
        "Word_Test.docx",
        "Word_Test_Ausgangsdokument.docx",
        "word_test_rheinfall.png",
    ),
}

_POINT_RE = re.compile(r"\[(\d+)\s*P\]")


def _docx_text(path: Path) -> str:
    with zipfile.ZipFile(path, "r") as archive:
        xml = archive.read("word/document.xml")
    root = ElementTree.fromstring(xml)
    return "".join(node.text or "" for node in root.iter() if node.tag.endswith("}t"))


def _validate_docx(path: Path) -> None:
    if path.stat().st_size < 1024:
        raise RuntimeError(f"DOCX is unexpectedly small: {path}")

    with zipfile.ZipFile(path, "r") as archive:
        bad_member = archive.testzip()
        if bad_member is not None:
            raise RuntimeError(f"Corrupt DOCX member {bad_member!r} in {path}")
        names = set(archive.namelist())
        required = {"[Content_Types].xml", "word/document.xml"}
        missing = sorted(required - names)
        if missing:
            raise RuntimeError(f"Invalid DOCX {path}; missing: {', '.join(missing)}")
        ElementTree.fromstring(archive.read("word/document.xml"))

    Document(path)


def _validate_png(path: Path) -> None:
    if path.stat().st_size == 0:
        raise RuntimeError(f"Empty PNG: {path}")
    with Image.open(path) as image:
        image.verify()
    with Image.open(path) as image:
        if image.width <= 0 or image.height <= 0:
            raise RuntimeError(f"Invalid PNG dimensions: {path}")


def _validate_student_package(path: Path, expected_members: tuple[str, ...]) -> None:
    if path.stat().st_size < 1024:
        raise RuntimeError(f"Student package is unexpectedly small: {path}")
    with zipfile.ZipFile(path, "r") as archive:
        bad_member = archive.testzip()
        if bad_member is not None:
            raise RuntimeError(f"Corrupt ZIP member {bad_member!r} in {path}")
        names = tuple(archive.namelist())
        if names != expected_members:
            raise RuntimeError(
                f"Unexpected contents in {path.name}: found {names}; expected {expected_members}."
            )
        for name in names:
            if Path(name).name != name:
                raise RuntimeError(f"Student package must be flat, but found nested member {name!r} in {path}")


def _validate_task_points(path: Path, expected_count: int, expected_sum: int) -> None:
    values = [int(value) for value in _POINT_RE.findall(_docx_text(path))]
    if len(values) != expected_count or sum(values) != expected_sum:
        raise RuntimeError(
            f"Unexpected points in {path.name}: found {values}; "
            f"expected {expected_count} tasks totalling {expected_sum}."
        )


def _validate_grade_rounding() -> None:
    expected = {
        (12, 20): "4.0",
        (13, 20): "4.3",
        (18, 30): "4.0",
        (30, 30): "6.0",
    }
    wrong = {
        key: (swiss_grade_str(*key), note)
        for key, note in expected.items()
        if swiss_grade_str(*key) != note
    }
    if wrong:
        raise RuntimeError(f"Swiss grade rounding is inconsistent: {wrong}")

    if str(effort_grade(26, 26)) != "6.0":
        raise RuntimeError("Full Fleisspunkte must produce grade 6.0.")

    final_examples = {
        ("5.3", "4.8", "3.5"): "5.1",
        ("4.2", "4.2", "5.0"): "4.6",
        ("6.0", "1.0", "6.0"): "6.0",
    }
    wrong_final = {
        grades: (final_grade_with_drop_str(*grades), note)
        for grades, note in final_examples.items()
        if final_grade_with_drop_str(*grades) != note
    }
    if wrong_final:
        raise RuntimeError(f"Final-grade drop rule is inconsistent: {wrong_final}")


def _validate_steckbrief(path: Path) -> None:
    doc = Document(path)
    rubric_rows = 0
    for table in doc.tables:
        for row in table.rows:
            if any(cell.text.strip() == "0 / 1 / 2" for cell in row.cells):
                rubric_rows += 1
    if rubric_rows != 10:
        raise RuntimeError(f"Steckbrief rubric has {rubric_rows} scored rows; expected 10.")

    text = _docx_text(path)
    if "maximal 20 Punkte" not in text or "12/20 = Note 4.0" not in text:
        raise RuntimeError("Steckbrief 20-point grading key is incomplete or inconsistent.")
    if "Kaufmännisch auf eine Dezimalstelle runden" not in text:
        raise RuntimeError("Steckbrief grading key does not define half-up rounding explicitly.")


def _validate_word_test_key(path: Path) -> None:
    doc = Document(path)
    mapping: dict[int, str] = {}
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            for point_col, note_col in ((0, 1), (2, 3)):
                if len(cells) <= note_col or not cells[point_col].isdigit():
                    continue
                mapping[int(cells[point_col])] = cells[note_col]

    if set(mapping) != set(range(31)):
        missing = sorted(set(range(31)) - set(mapping))
        raise RuntimeError(f"Word-test grading table does not cover 0-30 points; missing {missing}.")

    expected = {points: swiss_grade_str(points, 30) for points in range(31)}
    wrong = {
        points: (mapping.get(points), note)
        for points, note in expected.items()
        if mapping.get(points) != note
    }
    if wrong:
        raise RuntimeError(f"Word-test grading table is inconsistent with the shared grade function: {wrong}")

    text = _docx_text(path)
    if "Kaufmännisch auf eine Dezimalstelle runden" not in text:
        raise RuntimeError("Word-test grading key does not define half-up rounding explicitly.")


def validate_course_package(root: Path) -> tuple[int, int, int]:
    root = Path(root)
    output = root / "arbeitsblaetter"
    if not output.is_dir():
        raise RuntimeError(f"Generated output directory is missing: {output}")

    missing = [name for name in EXPECTED_DOCX if not (output / name).is_file()]
    missing += [name for name in EXPECTED_ASSETS if not (output / name).is_file()]
    missing += [name for name in EXPECTED_PACKAGES if not (output / name).is_file()]
    if not (output / "README.md").is_file():
        missing.append("README.md")
    if missing:
        raise RuntimeError("Build is incomplete; missing: " + ", ".join(sorted(missing)))

    docx_files = sorted(output.glob("*.docx"))
    for path in docx_files:
        _validate_docx(path)

    png_files = sorted((output / "assets").rglob("*.png"))
    if not png_files:
        raise RuntimeError("No generated PNG assets found.")
    for path in png_files:
        _validate_png(path)

    for package_rel, expected_members in EXPECTED_PACKAGES.items():
        _validate_student_package(output / package_rel, expected_members)

    _validate_grade_rounding()
    _validate_task_points(output / "Uebungstest_Word.docx", expected_count=10, expected_sum=30)
    _validate_task_points(output / "Word_Test.docx", expected_count=11, expected_sum=30)
    _validate_steckbrief(output / "Benoteter_Steckbrief.docx")
    _validate_word_test_key(output / "Word_Test_Korrekturblatt.docx")

    return len(docx_files), len(png_files), len(EXPECTED_PACKAGES)


def main() -> None:
    parser = argparse.ArgumentParser(description="Validate the generated Word course package.")
    parser.add_argument(
        "--root",
        type=Path,
        default=Path(__file__).resolve().parents[1],
        help="Repository or staged build root containing arbeitsblaetter/.",
    )
    args = parser.parse_args()
    docx_count, png_count, package_count = validate_course_package(args.root)
    print(
        f"Validated Word course package: {docx_count} DOCX, {png_count} PNG, "
        f"{package_count} student ZIP packages"
    )


if __name__ == "__main__":
    main()
