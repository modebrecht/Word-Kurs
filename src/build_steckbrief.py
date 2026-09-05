from __future__ import annotations

from pathlib import Path

from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

from course_common import (
    NAVY, TEAL_DARK, PALE_TEAL, WARM, MID, LIGHT_LINE,
    WHITE, TEXT, base_doc, block, add_text, finalise, set_run,
)
from course_build_helpers import (
    clear_paragraph as _clear,
    style_run as _font,
    fill_cell as _fill,
    set_cell_margins as _margins,
    set_cell_borders as _border,
)


def _req(cell, title: str, text: str):
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(.35)
    p.paragraph_format.line_spacing = 1
    r = p.add_run(title + ": ")
    set_run(r, size=9.25, bold=True, color=NAVY)
    r = p.add_run(text)
    set_run(r, size=9.25, color=TEXT)


def _custom_header_footer(doc):
    sec = doc.sections[0]
    ht = sec.header.tables[0]
    ht.cell(0,2).text = "BENOTETER AUFTRAG"
    for j, c in enumerate(ht.rows[0].cells):
        for p in c.paragraphs:
            for r in p.runs:
                _font(r, size=9.2, bold=True, color=WHITE if j == 0 else (TEAL_DARK if j == 2 else MID))
    ft = sec.footer.tables[0]
    ft.cell(0,1).text = "PERSÖNLICHER STECKBRIEF"
    for c in ft.rows[0].cells:
        for p in c.paragraphs:
            for r in p.runs:
                _font(r, size=8.2, bold=True, color=MID)


def build_steckbrief(out: Path):
    doc = base_doc(
        "NOTE",
        "Steckbrief",
        "Das bin ich",
        "Du erstellst einen persönlichen Steckbrief und zeigst dabei, was du in Word gelernt hast.",
        "Text, Absätze, Aufzählungen und ein Foto selbstständig, sauber und nach genauen Vorgaben gestalten.",
    )
    _custom_header_footer(doc)

    t = block(doc, "01", "INHALT")
    r = t.cell(0,1)
    p = r.paragraphs[0]
    _clear(p)
    x = p.add_run("Dein Steckbrief enthält alle diese Bereiche:")
    set_run(x, size=12.2, bold=True, color=NAVY)
    for text in [
        "dein Vorname als Haupttitel",
        "«Über mich» – 3 bis 5 ganze Sätze",
        "«Meine Interessen» – mindestens 3 Punkte",
        "«Das kann ich gut» – mindestens 2 Punkte",
        "«Das möchte ich noch lernen» – mindestens 1 Punkt",
        "«Mein Fun Fact» – 1 kurzer Satz",
        "ein eigenes bzw. von der Schule freigegebenes Foto",
    ]:
        add_text(r, "•  " + text, 9.2, after=.18)

    t = block(doc, "02", "FORMAT", fill_left=WARM, fill_right=WARM, label_color=NAVY)
    r = t.cell(0,1)
    _clear(r.paragraphs[0])
    _req(r, "Seite", "genau 1 Seite A4, Hochformat")
    _req(r, "Schrift", "Arial; Fliesstext 11 pt")
    _req(r, "Haupttitel", "20 pt, fett")
    _req(r, "Zwischentitel", "14 pt, fett")
    _req(r, "Absätze", "nach jedem Absatz 6 pt Abstand; keine Ketten aus Leerzeilen oder Leerzeichen")
    _req(r, "Liste", "«Meine Interessen» als echte Word-Aufzählung")

    t = block(doc, "03", "FOTO", fill_left=PALE_TEAL, fill_right=PALE_TEAL, label_color=TEAL_DARK)
    r = t.cell(0,1)
    _clear(r.paragraphs[0])
    _req(r, "Grösse", "ungefähr 5 cm breit")
    _req(r, "Bearbeitung", "sinnvoll zuschneiden; nicht verzerren")
    _req(r, "Textumbruch", "«Quadrat»")
    _req(r, "Position", "rechts oben neben «Über mich»; kein Text darf über dem Bild liegen")

    t = block(doc, "ABGABE", None, fill_left=WHITE, fill_right=WHITE, label_color=TEAL_DARK, label_size=9.2)
    r = t.cell(0,1)
    _clear(r.paragraphs[0])
    _req(r, "Dateiname", "Steckbrief_Nachname_Vorname.docx")
    _req(r, "Vor Abgabe", "genau 1 Seite, nichts abgeschnitten, kein Überlappen, alle Pflichtbereiche vorhanden")
    add_text(r, "Bewertet wird die korrekte Word-Arbeit – nicht, wie aussergewöhnlich deine Antworten sind.", 9.25, bold=True, color=TEAL_DARK, after=0)

    t = block(doc, "FOTO", None, fill_left=WHITE, fill_right=WHITE, label_color=TEAL_DARK, label_size=9.2)
    p = t.cell(0,1).paragraphs[0]
    _clear(p)
    x = p.add_run("Das Foto bleibt in der Schulabgabe. Lade den Steckbrief nicht öffentlich hoch.")
    set_run(x, size=9.35, color=MID)
    for c in t.rows[0].cells:
        _border(c, top={"val":"single","sz":"7","color":LIGHT_LINE})

    sec2 = doc.add_section(WD_SECTION.NEW_PAGE)
    sec2.page_width = Cm(21)
    sec2.page_height = Cm(29.7)
    sec2.top_margin = Cm(1.7)
    sec2.bottom_margin = Cm(1.5)
    sec2.left_margin = Cm(1.85)
    sec2.right_margin = Cm(1.85)
    sec2.header.is_linked_to_previous = True
    sec2.footer.is_linked_to_previous = True

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    x = p.add_run("BEWERTUNG  ·  STECKBRIEF")
    set_run(x, size=10, bold=True, color=TEAL_DARK)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    x = p.add_run("Bewertungsraster")
    set_run(x, size=26, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    x = p.add_run("10 Kriterien × 0 / 1 / 2 Punkte = maximal 20 Punkte")
    set_run(x, size=10.6, color=MID)

    rubric = [
        ("Pflichtinhalte", "Alle verlangten Bereiche und Mindestmengen sind vorhanden."),
        ("Seitenformat", "Genau 1 A4-Seite im Hochformat; nichts liegt ausserhalb der Seite."),
        ("Grundschrift", "Fliesstext durchgehend Arial 11 pt."),
        ("Haupttitel", "Vorname als Haupttitel, 20 pt und fett."),
        ("Zwischentitel", "Alle Bereichstitel 14 pt und fett; gleiche Stufe sieht gleich aus."),
        ("Absätze", "Saubere 6-pt-Abstände; keine Layout-Ketten aus Leerzeichen oder Leerzeilen."),
        ("Aufzählung", "«Meine Interessen» ist eine echte Word-Aufzählung."),
        ("Foto", "Foto sinnvoll zugeschnitten, ungefähr 5 cm breit und nicht verzerrt."),
        ("Bildplatzierung", "Textumbruch «Quadrat»; Foto rechts oben; kein Text überlappt."),
        ("Abgabe & Sauberkeit", "Dateiname korrekt; nichts abgeschnitten/überlappt; Dokument gut lesbar."),
    ]
    table = doc.add_table(rows=1, cols=5)
    table.autofit = False
    widths = [Cm(.9), Cm(4.0), Cm(8.2), Cm(1.6), Cm(1.6)]
    for i, (c, w, text) in enumerate(zip(table.rows[0].cells, widths, ["#","Kriterium","Erfüllt, wenn ...","P.","Notiz"])):
        c.width = w
        _fill(c, PALE_TEAL)
        _margins(c, 85, 100, 85, 100)
        p = c.paragraphs[0]
        _clear(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in (0,3,4) else WD_ALIGN_PARAGRAPH.LEFT
        x = p.add_run(text)
        set_run(x, size=9.2, bold=True, color=NAVY)
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    for idx, (crit, desc) in enumerate(rubric, 1):
        row = table.add_row()
        vals = [str(idx), crit, desc, "0 / 1 / 2", ""]
        for i, (c, w, text) in enumerate(zip(row.cells, widths, vals)):
            c.width = w
            _margins(c, 70, 90, 70, 90)
            c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = c.paragraphs[0]
            _clear(p)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in (0,3,4) else WD_ALIGN_PARAGRAPH.LEFT
            x = p.add_run(text)
            set_run(x, size=8.8, bold=(i == 1), color=NAVY if i == 1 else TEXT)
            _border(c,
                    top={"val":"single","sz":"4","color":LIGHT_LINE},
                    bottom={"val":"single","sz":"4","color":LIGHT_LINE},
                    left={"val":"single","sz":"4","color":LIGHT_LINE},
                    right={"val":"single","sz":"4","color":LIGHT_LINE})

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    x = p.add_run("Punkte: ____ / 20     Note: ____")
    set_run(x, size=12, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    x = p.add_run("0 = fehlt / klar falsch   ·   1 = teilweise erfüllt   ·   2 = vollständig erfüllt")
    set_run(x, size=9.3, color=MID)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    x = p.add_run("Notenschlüssel: Note = 1 + 5 × (Punkte / 20). Kaufmännisch auf eine Dezimalstelle runden. 12/20 = Note 4.0.")
    set_run(x, size=9.3, bold=True, color=TEAL_DARK)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    x = p.add_run("Kreativität oder «schönes Design» ist kein eigenes Bewertungskriterium.")
    set_run(x, size=9.3, color=MID)

    finalise(doc, out, "Benoteter persönlicher Steckbrief")


def build_all(root: Path):
    build_steckbrief(root / "arbeitsblaetter" / "Benoteter_Steckbrief.docx")
