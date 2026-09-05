from __future__ import annotations

from docx.shared import Cm, Pt
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from course_common import (
    TEAL_DARK,
    set_run,
    _border,
    _clear,
    _fill,
    _font,
    _margins,
)


def clear_paragraph(paragraph):
    """Clear paragraph content while preserving paragraph properties."""
    return _clear(paragraph)


def style_run(run, *, name="Arial", size=None, bold=None, italic=None, color=None):
    """Apply the course's explicit Word run formatting."""
    return _font(run, name=name, size=size, bold=bold, italic=italic, color=color)


def fill_cell(cell, color):
    """Set a Word table cell fill color."""
    return _fill(cell, color)


def set_cell_margins(cell, top=80, start=145, bottom=80, end=145):
    """Set Word table cell margins in twips."""
    return _margins(cell, top, start, bottom, end)


def set_cell_borders(cell, **edges):
    """Set Word table cell border properties."""
    return _border(cell, **edges)


def restart_page_numbering(section, start=1):
    """Restart automatic page numbering for a section."""
    sect_pr = section._sectPr
    page_number = sect_pr.find(qn("w:pgNumType"))
    if page_number is None:
        page_number = OxmlElement("w:pgNumType")
        sect_pr.append(page_number)
    page_number.set(qn("w:start"), str(start))


def new_detached_workspace_section(
    doc,
    code: str | None = None,
    *,
    top_margin_cm: float = 1.8,
    bottom_margin_cm: float = 1.8,
    left_margin_cm: float = 2.0,
    right_margin_cm: float = 2.0,
    header_distance_cm: float | None = None,
    footer_distance_cm: float | None = None,
    page_number_start: int = 1,
):
    """Create a clean A4 workspace section with independent header/footer.

    When ``code`` is provided, the section starts with the standard
    ``<code> · ARBEITSSEITE`` label. Header/footer tables inherited from the
    previous section are removed so the workspace is genuinely blank.
    """
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(top_margin_cm)
    section.bottom_margin = Cm(bottom_margin_cm)
    section.left_margin = Cm(left_margin_cm)
    section.right_margin = Cm(right_margin_cm)
    if header_distance_cm is not None:
        section.header_distance = Cm(header_distance_cm)
    if footer_distance_cm is not None:
        section.footer_distance = Cm(footer_distance_cm)

    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    for part in (section.header, section.footer):
        for paragraph in part.paragraphs:
            clear_paragraph(paragraph)
        for table in list(part.tables):
            table._tbl.getparent().remove(table._tbl)

    restart_page_numbering(section, page_number_start)

    if code:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(5)
        run = paragraph.add_run(f"{code} · ARBEITSSEITE")
        set_run(run, size=9.3, bold=True, color=TEAL_DARK)

    return section
