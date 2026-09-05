from __future__ import annotations

from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK

from course_common import (
    NAVY, TEAL, TEAL_DARK, PALE_TEAL, WARM, MID, TEXT,
    base_doc, block, add_text, add_step, add_tip, add_check, add_finish,
    add_picture, finalise, set_run,
)
from course_build_helpers import (
    clear_paragraph as _clear,
    style_run as _font,
    new_detached_workspace_section,
)


def _fonts():
    root = Path('/usr/share/fonts/truetype/liberation2')
    return str(root/'LiberationSans-Regular.ttf'), str(root/'LiberationSans-Bold.ttf')


def a10_preview(path: Path):
    reg, bold = _fonts()
    W, H = 1500, 440
    im = Image.new('RGB', (W, H), 'white')
    d = ImageDraw.Draw(im)
    f_small = ImageFont.truetype(reg, 24)
    f_small_b = ImageFont.truetype(bold, 24)
    f_num = ImageFont.truetype(bold, 25)
    line = '#D3DEE2'

    d.rounded_rectangle((18, 18, W-18, H-18), radius=18, outline=line, width=3, fill='white')

    # Two mini pages show repetition across pages without UI screenshots.
    for idx, x in enumerate((120, 800), start=1):
        y = 54
        pw, ph = 520, 330
        d.rectangle((x, y, x+pw, y+ph), fill='white', outline='#9FB0B7', width=3)
        d.line((x+30, y+55, x+pw-30, y+55), fill=line, width=2)
        d.line((x+30, y+ph-52, x+pw-30, y+ph-52), fill=line, width=2)
        d.text((x+34, y+17), 'SCHULAUSFLUG LUZERN', font=f_small_b, fill='#17324D')
        d.text((x+40, y+92), 'Reisebericht', font=f_small_b, fill='#237B78')
        d.text((x+40, y+135), 'Der Text steht im normalen Seitenbereich.', font=f_small, fill='#5E6D78')
        footer = f'Seite {idx}'
        bbox = d.textbbox((0, 0), footer, font=f_num)
        d.text((x+(pw-(bbox[2]-bbox[0]))/2, y+ph-42), footer, font=f_num, fill='#17324D')

    im.save(path, quality=95)


def build_a10(out: Path, preview: Path):
    doc = base_doc(
        'A10', 'Kopf- & Fusszeile', 'Infos auf jeder Seite',
        'Du setzt Informationen in den oberen und unteren Seitenbereich.',
        'eine Kopfzeile, eine Fusszeile und automatische Seitenzahlen einfügen.'
    )

    t = block(doc, '01', 'SEITEN 2–3')
    r = t.cell(0,1)
    p = r.paragraphs[0]
    _clear(p)
    x = p.add_run('Ergänze den Reisebericht')
    set_run(x, size=12.8, bold=True, color=NAVY)
    add_text(r, 'Arbeite nur auf den Seiten 2 und 3. Der normale Text ist bereits fertig.', 9.5, after=.6)
    add_picture(r, preview, 10.9)

    add_step(r, 'A', [('Doppelklicke ganz oben auf Seite 2', True, False, NAVY), ('  →  Kopfzeile öffnen', False, False, None)])
    add_step(r, 'B', [('Kopfzeile', True, False, NAVY), ('  →  «SCHULAUSFLUG LUZERN» eingeben', False, False, None)])
    add_step(r, 'C', [('Doppelklicke ganz unten auf Seite 2', True, False, NAVY), ('  →  Fusszeile öffnen', False, False, None)])
    add_step(r, 'D', [('Fusszeile', True, False, NAVY), ('  →  «Seite » eingeben', False, False, None)])
    add_step(r, 'E', [('Direkt hinter «Seite »', True, False, NAVY), ('  →  automatische Seitenzahl einfügen', False, False, None)])
    add_step(r, 'F', [('Kontrolliere Seite 3', True, False, NAVY), ('  →  Kopf- und Fusszeile erscheinen dort automatisch.', False, False, None)])

    add_tip(doc, 'Kopf- und Fusszeile sind eigene Seitenbereiche. Tippe diese Angaben nicht in den normalen Text.')
    add_tip(doc, 'Für die Zahl: Einfügen → Seitenzahl → Aktuelle Position. Die Nummer muss automatisch sein – nicht von Hand tippen.', 'MERKE')
    add_check(doc, 'Oben steht auf beiden Übungsseiten «SCHULAUSFLUG LUZERN». Unten steht «Seite 1» bzw. «Seite 2».')
    add_finish(doc)

    # Workspace section: deliberately no course header/footer. Page numbering is already
    # configured to restart at 1, so students only learn the intended skill: insert PAGE.
    new_detached_workspace_section(
        doc,
        top_margin_cm=2.2,
        bottom_margin_cm=2.0,
        left_margin_cm=2.0,
        right_margin_cm=2.0,
        header_distance_cm=.7,
        footer_distance_cm=.7,
    )

    # Page 2 content.
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    x = p.add_run('REISEBERICHT LUZERN')
    set_run(x, size=20, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    x = p.add_run('Der Morgen')
    set_run(x, size=15, bold=True, color=TEAL)

    for s in [
        'Um 08.00 Uhr treffen wir uns beim Schulhaus. Gemeinsam fahren wir mit dem Zug nach Luzern.',
        'Nach der Ankunft gehen wir direkt zum Verkehrshaus. Dort arbeiten wir in kleinen Gruppen.',
        'Jede Gruppe sucht drei Ausstellungsstücke und notiert dazu die wichtigsten Informationen.',
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        x = p.add_run(s)
        set_run(x, size=11)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(5)
    x = p.add_run('Im Verkehrshaus')
    set_run(x, size=15, bold=True, color=TEAL)

    for s in [
        'Besonders spannend sind die Bereiche Luftfahrt und Raumfahrt. Einige Gruppen besuchen zusätzlich die Eisenbahn-Ausstellung.',
        'Vor dem Mittag treffen wir uns wieder beim Haupteingang und vergleichen unsere Notizen.',
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        x = p.add_run(s)
        set_run(x, size=11)

    # Page 3, same section: the student's header/footer should repeat automatically.
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)  # normal page break, same section

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    x = p.add_run('Mittag und Altstadt')
    set_run(x, size=15, bold=True, color=TEAL)

    for s in [
        'Das Mittagessen verbringen wir am See. Danach laufen wir gemeinsam in Richtung Altstadt.',
        'Beim Rundgang sehen wir die Kapellbrücke und verschiedene historische Gebäude. Anschliessend bleibt Zeit für eine kurze Pause.',
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        x = p.add_run(s)
        set_run(x, size=11)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(5)
    x = p.add_run('Rückfahrt')
    set_run(x, size=15, bold=True, color=TEAL)

    for s in [
        'Um 16.30 Uhr nehmen wir den Zug zurück. Die Ankunft beim Schulhaus ist ungefähr um 18.00 Uhr.',
        'Damit endet unser Schulausflug nach Luzern.',
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        x = p.add_run(s)
        set_run(x, size=11)

    finalise(doc, out, 'A10 - Kopf- und Fusszeile / Seitenzahlen')


def a11_asset(path: Path):
    W,H=1200,800
    im=Image.new('RGB',(W,H),'#F7FBFC'); d=ImageDraw.Draw(im)
    d.rectangle((0,0,W,220),fill='#DCECF5'); d.rectangle((0,220,W,520),fill='#EAF4F8'); d.rectangle((0,520,W,H),fill='#DCE9E3')
    d.ellipse((110,540,1090,770),fill='#A9D2D5'); d.ellipse((180,585,1020,720),fill='#96C3C7')
    d.polygon([(70,565),(370,185),(675,565)],fill='#6E828E'); d.polygon([(390,565),(735,115),(1100,565)],fill='#566C78')
    d.polygon([(245,355),(370,185),(500,355)],fill='white'); d.polygon([(585,305),(735,115),(900,315)],fill='white')
    d.line([(70,565),(370,185),(675,565)],fill='#4F6672',width=4); d.line([(390,565),(735,115),(1100,565)],fill='#415863',width=4)
    for x,y,scale in [(125,495,1.0),(1015,485,1.08),(935,525,.82),(235,535,.76)]:
        d.rectangle((x-8,y+35,x+8,y+102),fill='#7A5B43')
        d.polygon([(x,y-50*scale),(x-55*scale,y+52*scale),(x+55*scale,y+52*scale)],fill='#4A7568')
        d.polygon([(x,y-5*scale),(x-48*scale,y+75*scale),(x+48*scale,y+75*scale)],fill='#3D6257')
    d.rectangle((4,4,W-5,H-5),outline='#D3DEE2',width=8); im.save(path,quality=95)


def a11_preview(path: Path, asset: Path):
    reg,bold=_fonts(); W,H=1600,980
    im=Image.new('RGB',(W,H),'white'); d=ImageDraw.Draw(im)
    fh=ImageFont.truetype(bold,22); ft=ImageFont.truetype(bold,52); f1=ImageFont.truetype(bold,30); fb=ImageFont.truetype(reg,24); fbb=ImageFont.truetype(bold,24); fs=ImageFont.truetype(reg,20); ftab=ImageFont.truetype(reg,21); ftabb=ImageFont.truetype(bold,21)
    d.rounded_rectangle((18,18,W-18,H-18),radius=18,outline='#D3DEE2',width=3,fill='white')
    d.text((65,45),'KLASSENLAGER 2026',font=fh,fill='#17324D'); d.text((1320,45),'SEK 8B',font=fh,fill='#237B78'); d.line((65,85,1535,85),fill='#237B78',width=4)
    d.text((65,125),'KLASSENLAGER FLIMS',font=ft,fill='#17324D'); d.text((67,190),'15.–18. Juni 2026',font=fbb,fill='#237B78')
    pic=Image.open(asset).convert('RGB').crop((110,80,1090,750)); pic.thumbnail((520,340)); im.paste(pic,(1000,135)); d.rectangle((1000,135,1000+pic.width,135+pic.height),outline='#D3DEE2',width=2)
    d.text((65,255),'Vier Tage unterwegs in Graubünden:',font=fb,fill='#17324D'); d.text((65,295),'wandern, gemeinsam kochen und Zeit am See.',font=fb,fill='#17324D')
    d.text((65,385),'MITNEHMEN',font=f1,fill='#237B78'); y=435
    for item in ['Wanderschuhe','Regenjacke','Trinkflasche','kleiner Rucksack']:
        d.ellipse((74,y+8,84,y+18),fill='#17324D'); d.text((100,y),item,font=fb,fill='#17324D'); y+=42
    d.text((65,625),'PROGRAMM',font=f1,fill='#237B78')
    x0,y0=65,675; cols=[210,390,410]; rows=[['Tag','Vormittag','Nachmittag'],['Montag','Anreise','Dorfrundgang'],['Dienstag','Wanderung','Caumasee'],['Mittwoch','Sport','Freizeit'],['Donnerstag','Aufräumen','Rückreise']]
    y=y0
    for ri,row in enumerate(rows):
        x=x0
        for w,val in zip(cols,row):
            d.rectangle((x,y,x+w,y+48 if ri==0 else y+45),fill='#EAF4F3' if ri==0 else 'white',outline='#D3DEE2',width=2); d.text((x+12,y+12 if ri==0 else y+11),val,font=ftabb if ri==0 else ftab,fill='#17324D'); x+=w
        y+=48 if ri==0 else 45
    d.line((65,935,1535,935),fill='#D3DEE2',width=2); d.text((65,947),'Klassenlager Flims',font=fs,fill='#667684'); d.text((1440,947),'1',font=fs,fill='#667684')
    im.save(path,quality=95)


def build_a11(out: Path, preview: Path):
    doc=base_doc('A11','Dokument nachbauen','Jetzt ohne Klickanleitung','Du kombinierst bekannte Word-Werkzeuge und baust eine Vorlage möglichst genau nach.','mehrere bekannte Word-Funktionen selbstständig kombinieren und ein Dokument nach einer sichtbaren Vorlage nachbauen.')
    t=block(doc,'01','NACHBAUEN'); r=t.cell(0,1); p=r.paragraphs[0]; _clear(p); x=p.add_run('Baue Seite 2 so um, dass sie wie diese Vorlage aussieht'); set_run(x,size=12.6,bold=True,color=NAVY); add_text(r,'Der Text und die Daten sind vorgegeben. Du entscheidest selbst, welche bekannten Werkzeuge du dafür brauchst.',9.35); add_picture(r,preview,10.9)
    t=block(doc,'PFLICHT',None,fill_left=WARM,fill_right=WARM,label_color=NAVY,label_size=9.3); p=t.cell(0,1).paragraphs[0]; _clear(p); x=p.add_run('Verwende wirklich: '); set_run(x,size=9.55,bold=True,color=NAVY); x=p.add_run('Formatvorlagen für Titel/Überschriften · echte Aufzählung · echte Tabelle · Bild mit Textumbruch · Kopfzeile · automatische Seitenzahl.'); set_run(x,size=9.25)
    add_tip(doc,'Arbeite von gross nach klein: zuerst Aufbau und Bereiche, danach Bild/Tabelle, ganz am Schluss Abstände und Feinarbeit.')
    add_check(doc,'Stimmen Reihenfolge, Grössenverhältnisse, Bildposition, Liste, Tabelle sowie Kopf- und Fussbereich?')
    add_finish(doc,'Gib dieses Arbeitsblatt zusammen mit der Bilddatei in deinem Ordner "IB" ab.')
    new_detached_workspace_section(doc,'A11',top_margin_cm=1.8,bottom_margin_cm=1.8)
    for text in ['KLASSENLAGER FLIMS','15.–18. Juni 2026','Vier Tage unterwegs in Graubünden: wandern, gemeinsam kochen und Zeit am See.','MITNEHMEN','Wanderschuhe','Regenjacke','Trinkflasche','kleiner Rucksack','PROGRAMM','Tag | Vormittag | Nachmittag','Montag | Anreise | Dorfrundgang','Dienstag | Wanderung | Caumasee','Mittwoch | Sport | Freizeit','Donnerstag | Aufräumen | Rückreise']:
        p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(3.5); x=p.add_run(text); set_run(x,size=11)
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(8); x=p.add_run('BILDDATEI: a11_klassenlager_berge.png'); set_run(x,size=9.5,bold=True,color=TEAL_DARK)
    finalise(doc,out,'A11 - Dokument nach Vorlage nachbauen')


def a12_asset(path: Path):
    W,H=1600,900; im=Image.new('RGB',(W,H),'#E7F1F4'); d=ImageDraw.Draw(im)
    d.rectangle((0,0,W,330),fill='#D9E9EF'); d.rectangle((0,330,W,610),fill='#F1E6D4'); d.rectangle((0,610,W,H),fill='#BED5C7')
    d.rectangle((360,310,1240,690),fill='#F4F0E7',outline='#17324D',width=7); d.polygon([(310,320),(800,95),(1290,320)],fill='#667D88',outline='#17324D')
    d.rectangle((720,500,880,690),fill='#B8D4D7',outline='#17324D',width=5)
    for y in (380,500):
        for x in (440,570,960,1090): d.rectangle((x,y,x+90,y+75),fill='#D8EEF1',outline='#17324D',width=4)
    for x in (180,1370): d.rectangle((x-15,480,x+15,700),fill='#7A5B43'); d.ellipse((x-115,330,x+115,560),fill='#4D7568')
    d.line((130,220,1470,220),fill='#445B65',width=4)
    for x in range(180,1450,120): d.line((x,220,x,252),fill='#445B65',width=3); d.ellipse((x-10,245,x+10,265),fill='#F5C86B',outline='#A67B31')
    for x in (270,1050): d.rectangle((x,710,x+250,742),fill='#8A6549'); d.line((x+40,742,x+15,800),fill='#8A6549',width=12); d.line((x+210,742,x+235,800),fill='#8A6549',width=12); d.rectangle((x-20,770,x+270,792),fill='#8A6549')
    d.rectangle((4,4,W-5,H-5),outline='#D3DEE2',width=8); im.save(path,quality=95)


def build_a12(out: Path):
    doc=base_doc('A12','Selbstständig gestalten','Du entscheidest','Es gibt keine Zielvorlage mehr. Die Vorgaben sind klar – das Layout planst du selbst.','ein einseitiges Dokument mit bekannten Word-Werkzeugen selbstständig planen und übersichtlich gestalten.')
    t=block(doc,'01','FLYER'); r=t.cell(0,1); p=r.paragraphs[0]; _clear(p); x=p.add_run('Gestalte auf Seite 2 einen Flyer für den Schulhaus-Sommerabend'); set_run(x,size=12.5,bold=True,color=NAVY); add_text(r,'Alle Informationen stehen bereit. Du darfst die Reihenfolge der Bereiche verändern, aber keinen Inhalt weglassen.',9.35)
    t=block(doc,'PFLICHT',None,fill_left=WARM,fill_right=WARM,label_color=NAVY,label_size=9.3); r=t.cell(0,1); p=r.paragraphs[0]; _clear(p); x=p.add_run('Dein Flyer muss enthalten:'); set_run(x,size=9.7,bold=True,color=NAVY)
    for s in ['A4 Hochformat · alles auf einer Seite','Formatvorlagen für Titel und Überschriften','eine echte Aufzählung bei «MITBRINGEN»','eine echte Tabelle für das Programm','das Bild a12_sommerabend.png · zugeschnitten und sinnvoll platziert','Fusszeile «Sommerabend 2026» + automatische Seitenzahl']:
        add_text(r,'•  '+s,9.3,after=.25)
    t=block(doc,'DU','ENTSCHEIDEST',fill_left=PALE_TEAL,fill_right=PALE_TEAL,label_color=TEAL_DARK,label_size=13.5); p=t.cell(0,1).paragraphs[0]; _clear(p); x=p.add_run('Du entscheidest selbst über Anordnung, Grössen, Abstände, Bildposition und eine passende Farbe. '); set_run(x,size=9.35); x=p.add_run('Der Flyer soll vor allem schnell lesbar sein.'); set_run(x,size=9.35,bold=True,color=TEAL_DARK)
    add_check(doc,'Sind Datum, Zeit und Ort sofort sichtbar? Ist alles auf einer Seite? Sind Liste, Tabelle, Bild und Fusszeile wirklich vorhanden?')
    add_finish(doc,'Gib dieses Arbeitsblatt zusammen mit der Bilddatei in deinem Ordner "IB" ab.')
    new_detached_workspace_section(doc,'A12',top_margin_cm=1.8,bottom_margin_cm=1.8)
    raw=['SOMMERABEND IM SCHULHAUS','Freitag, 19. Juni 2026','18.00–21.00 Uhr','Schulhaus Sonnenberg · Innenhof','Wir beenden das Schuljahr gemeinsam mit Musik, Spielen, Essen und Zeit zum Zusammensitzen.','PROGRAMM','18.00 | Start und Musik','18.45 | Grill und Getränke','19.30 | Spielturnier','20.30 | Dessert und Abschluss','MITBRINGEN','Trinkbecher','Jacke oder Pullover','gute Laune','wer möchte: ein Kartenspiel','WICHTIG','Bei Regen findet der Sommerabend in der Aula statt. Die Teilnahme ist kostenlos.','BILDDATEI: a12_sommerabend.png']
    for text in raw:
        p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(4); x=p.add_run(text); set_run(x,size=11)
    finalise(doc,out,'A12 - Selbstständig gestalten')


def build_all(root: Path):
    sheets = root/'arbeitsblaetter'
    assets = sheets/'assets'
    prev = assets/'vorlagen'
    assets.mkdir(parents=True, exist_ok=True)
    prev.mkdir(parents=True, exist_ok=True)

    a10prev = prev/'a10_kopf_fuss_vorlage.png'
    a10_preview(a10prev)
    build_a10(sheets/'A10_Kopf_Fusszeile_Seitenzahlen.docx', a10prev)

    a11img = assets/'a11_klassenlager_berge.png'
    a11prev = prev/'a11_klassenlager_vorlage.png'
    a11_asset(a11img)
    a11_preview(a11prev, a11img)
    build_a11(sheets/'A11_Dokument_nach_Vorlage_nachbauen.docx', a11prev)

    a12img = assets/'a12_sommerabend.png'
    a12_asset(a12img)
    build_a12(sheets/'A12_Selbststaendig_gestalten.docx')
